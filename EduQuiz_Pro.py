import tkinter as tk
from tkinter import messagebox, filedialog, ttk
import openpyxl
import os
import docx
import random  # Thêm import random
import shutil
from tkinter import scrolledtext
from docx import Document
import subprocess

# QUẢN LÍ TÀI KHOẢN
# Đường dẫn tới file Excel lưu danh sách tài khoản
accounts_file = "accounts.xlsx"

# Hàm khởi tạo file Excel nếu chưa tồn tại
def create_accounts_file():
    if not os.path.exists(accounts_file):
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Accounts"
        sheet.append(["Role", "Username", "Password"])  # Tiêu đề cột
        workbook.save(accounts_file)

# Hàm tạo file mẫu Excel để người dùng tải về
def download_sample_file():
    sample_file = "sample_accounts.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Accounts"
    sheet.append(["Role", "Username", "Password"])  # Cấu trúc mẫu
    sheet.append(["Teacher", "VuongVanHien", "Hien123"])
    sheet.append(["Student", "Nguyễn Văn Giàu", "giau456"])
    workbook.save(sample_file)
    
    # Hộp thoại lưu file
    file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
    if file_path:
        shutil.copy(sample_file, file_path)
        messagebox.showinfo("Thành công", "File mẫu đã được tải về!")

# Hàm nhập tài khoản từ file Excel
def import_accounts():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    if file_path:
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active
        workbook_accounts = openpyxl.load_workbook(accounts_file)
        sheet_accounts = workbook_accounts["Accounts"]
        
        for row in sheet.iter_rows(min_row=2, values_only=True):  # Bỏ qua hàng tiêu đề
            role, username, password = row
            if username and password:
                # Kiểm tra trùng lặp
                exists = False
                for existing_row in sheet_accounts.iter_rows(min_row=2, values_only=True):
                    if existing_row[1] == username:
                        exists = True
                        break
                if not exists:
                    sheet_accounts.append([role, username, password])
        
        workbook_accounts.save(accounts_file)
        messagebox.showinfo("Thành công", "Nhập tài khoản thành công!")

# CẤP TÀI KHOẢN
def add_account():
    def handle_add():
        role = role_var.get()
        username = username_entry.get().strip()
        password = password_entry.get().strip()

        if username and password:
            workbook = openpyxl.load_workbook(accounts_file)
            sheet = workbook["Accounts"]
            for row in sheet.iter_rows(min_row=2, values_only=True):
                if row[1] == username:
                    messagebox.showerror("Lỗi", "Tên đăng nhập đã tồn tại!")
                    return
            
            sheet.append([role, username, password])
            workbook.save(accounts_file)
            messagebox.showinfo("Thành công", "Tạo tài khoản thành công!")
            account_frame.forget()
        else:
            messagebox.showerror("Lỗi", "Vui lòng nhập tên đăng nhập và mật khẩu!")
        
    global change_frame  # Định nghĩa change_frame là biến toàn cục
    for widget in root.winfo_children():
        widget.forget()
    account_frame = tk.Frame(root, bg="#e6f2ff")  # Khởi tạo account_frame trong root
    account_frame.pack(fill="both", expand=True)
  

    tk.Label(account_frame, text="CẤP TÀI KHOẢN CHO NGƯỜI DÙNG", font=("Times New Roman", 18, 'bold'), bg="#e6f2ff", fg="#003366").pack(pady=20)
    tk.Label(account_frame, text="Vai trò:", font=("Times New Roman", 16, 'bold'), bg="#e6f2ff", fg="#FF0000").pack(pady=10)

    # Căn giữa hai Radiobutton
    global role_var
    role_var = tk.StringVar(value="Teacher")
    radio_frame = tk.Frame(account_frame, bg="#f0f0f0")
    radio_frame.pack(pady=10)
    
    tk.Radiobutton(radio_frame, text="Giáo viên", variable=role_var, value="Teacher", font=("Times New Roman", 14, 'bold'), bg="#e6f2ff", fg="#0000FF").pack(side=tk.TOP, anchor=tk.CENTER)
    tk.Radiobutton(radio_frame, text="Học sinh", variable=role_var, value="Student", font=("Times New Roman", 14, 'bold'), bg="#e6f2ff", fg="#0000FF").pack(side=tk.TOP, anchor=tk.CENTER)
    global username_entry
    tk.Label(account_frame, text="Tên đăng nhập:", font=("Times New Roman", 14, 'bold'), bg="#e6f2ff", fg="#003366").pack(pady=10)
    username_entry = tk.Entry(account_frame, width=30, font=("Times New Roman", 12))
    username_entry.pack(pady=5)
    global password_entry
    tk.Label(account_frame, text="Mật khẩu:", font=("Times New Roman", 14, 'bold'), bg="#f0f0f0", fg="#003366").pack(pady=10)
    password_entry = tk.Entry(account_frame, width=30, font=("Times New Roman", 12), show="*")
    password_entry.pack(pady=5)

    tk.Button(account_frame, text="Thêm", command=handle_add, font=("Times New Roman", 14, 'bold'), bg="#003366", fg="white").pack(pady=20)
    tk.Button(account_frame, text="Thoát", command=account_frame.forget, font=("Times New Roman", 14, 'bold'), bg="#800000", fg="white").pack(pady=10)



# Hàm kiểm tra thông tin đăng nhập
def check_login(role, username, password):
    workbook = openpyxl.load_workbook(accounts_file)
    sheet = workbook["Accounts"]
    for row in sheet.iter_rows(min_row=2, values_only=True):  # Bỏ qua hàng tiêu đề
        if row[0] == role and row[1] == username and row[2] == password:
            return True
    return False
# Hàm xử lý khi đăng nhập
def handle_login(username_entry, password_entry, role_var, login_frame):
    role = role_var.get()
    username = username_entry.get().strip()
    password = password_entry.get().strip()

    if check_login(role, username, password):
        messagebox.showinfo("Thành công", "Đăng nhập thành công!")
        login_frame.destroy()
        show_menus(role)  # Hiển thị menu dựa trên vai trò đăng nhập
    else:
        messagebox.showerror("Lỗi", "Tên đăng nhập hoặc mật khẩu không đúng!")

# ĐĂNG NHẬP HỆ THỐNG
# Hàm xử lý khi giáo viên hoặc học sinh đăng nhập
def user_login(role):
    login_frame = tk.Frame(root, bg="#e6f2ff")
    login_frame.pack(fill="both", expand=True)
    
    tk.Label(login_frame, text="ĐĂNG NHẬP HỆ THỐNG", font=("Times New Roman", 18, 'bold'), bg="#e6f2ff", fg="#003366").pack(pady=20)
    tk.Label(login_frame, text="Vai trò:", font=("Times New Roman", 16, 'bold'), bg="#e6f2ff", fg="#FF0000").pack(pady=5)

    # Biến toàn cục cho role
    role_var = tk.StringVar(value=role)

    # Tạo giao diện cho các nút radio
    radio_frame = tk.Frame(login_frame, bg="#e6f2ff")
    radio_frame.pack(pady=5)

    tk.Radiobutton(radio_frame, text="Giáo viên", variable=role_var, value="Teacher", font=("Times New Roman", 14, 'bold'), bg="#e6f2ff", fg="#0000FF").pack(side=tk.TOP, anchor=tk.CENTER)
    tk.Radiobutton(radio_frame, text="Học sinh", variable=role_var, value="Student", font=("Times New Roman", 14, 'bold'), bg="#e6f2ff", fg="#0000FF").pack(side=tk.TOP, anchor=tk.CENTER)

    # Nhập tên đăng nhập
    tk.Label(login_frame, text="Tên đăng nhập:", font=("Times New Roman", 14, 'bold'), bg="#e6f2ff", fg="#003366").pack(pady=10)
    username_entry = tk.Entry(login_frame, width=30, font=("Times New Roman", 12))
    username_entry.pack(pady=5)

    # Nhập mật khẩu
    tk.Label(login_frame, text="Mật khẩu:", font=("Times New Roman", 14, 'bold'), bg="#e6f2ff", fg="#003366").pack(pady=10)
    password_entry = tk.Entry(login_frame, width=30, font=("Times New Roman", 12), show="*")
    password_entry.pack(pady=5)

    # Nút Đăng nhập
    tk.Button(login_frame, text="Đăng nhập", command=lambda: handle_login(username_entry, password_entry, role_var, login_frame), font=("Times New Roman", 14, 'bold'), bg="#003366", fg="white").pack(pady=20)
    # Nút thoát
    tk.Button(login_frame, text="Thoát", command=login_frame.destroy, font=("Times New Roman", 14, 'bold'), bg="#800000", fg="white").pack(pady=10)

# Hàm ẩn các menu chưa đăng nhập
def hide_menus():
    menubar.entryconfig("Admin", state="disabled")
    menubar.entryconfig("Giáo viên", state="disabled")
    menubar.entryconfig("Kiểm tra", state="disabled")

# Hàm hiển thị các menu dựa trên vai trò
def show_menus(role):
    menubar.entryconfig("Học sinh", state="normal")
    menubar.entryconfig("Kiểm tra", state="normal")
    if role == "Teacher":
        menubar.entryconfig("Giáo viên", state="normal")
        menubar.entryconfig("Admin", state="normal")
    else:
        menubar.entryconfig("Giáo viên", state="disabled")
        menubar.entryconfig("Admin", state="disabled")
        
# Thông báo chức năng đang được nâng cấp
def teacher_menu_theory_lesson():
    messagebox.showinfo("Thông báo", "Chức năng đang nâng cấp")

# Tính toán để căn giữa cửa sổ
def center_window(window_width, window_height, window_root):
    screen_width = window_root.winfo_screenwidth()
    screen_height = window_root.winfo_screenheight()
    position_top = int((screen_height - window_height) / 2)
    position_right = int((screen_width - window_width) / 2)
    # Đặt vị trí cửa sổ
    window_root.geometry(f"{window_width}x{window_height}+{position_right}+{position_top}")
    
# Trật tự cửa sổ
def window_order(window_root):
    def on_close():
        window_root.grab_release()
        window_root.destroy()
    window_root.protocol("WM_DELETE_WINDOW", on_close)

# THAY ĐỔI TÀI KHOẢN ĐĂNG NHẬP        
# Hàm xóa tài khoản cũ và thêm tài khoản mới
def delete_and_update_account(role, username, password, new_username, new_password):
    try:
        workbook = openpyxl.load_workbook(accounts_file)
        sheet = workbook["Accounts"]
        deleted = False
        # Xoá tài khoản cũ
        for row in sheet.iter_rows(min_row=2):
            if row[0].value == role and row[1].value == username and row[2].value == password:
                sheet.delete_rows(row[0].row, 1)
                deleted = True
                break

        # Thêm tài khoản mới nếu xoá thành công
        if deleted:
            sheet.append([role, new_username, new_password])
            workbook.save(accounts_file)
            messagebox.showinfo("Thành công", "Tài khoản đã được cập nhật mới!")
        else:
            messagebox.showerror("Lỗi", "Không tìm thấy tài khoản để cũ.")
    except PermissionError:
        messagebox.showerror("Lỗi", f"Không thể ghi vào tệp '{accounts_file}'. Hãy chắc chắn tệp không đang được mở và bạn có quyền truy cập.")
    except Exception as e:
        messagebox.showerror("Lỗi", f"Đã xảy ra lỗi: {str(e)}")

# Hàm xử lý giao diện xoá và cập nhật tài khoản
def handle_change():
    role = role_var.get()
    username = username_entry.get().strip()
    password = password_entry.get().strip()
    new_username = new_username_entry.get().strip()
    new_password = new_password_entry.get().strip()

    if check_login(role, username, password):
        delete_and_update_account(role, username, password, new_username, new_password)
    else:
        messagebox.showerror("Lỗi", "Thông tin tài khoản cũ không đúng!")

def change_account(role):
    global change_frame  # Định nghĩa change_frame là biến toàn cục
    if 'change_frame' not in globals():  # Kiểm tra xem change_frame đã được định nghĩa chưa
        change_frame = tk.Frame(root, bg="#e6f2ff")
        change_frame.pack(fill="both", expand=True)

        # Tạo giao diện cho change_frame
        tk.Label(change_frame, text="THAY ĐỔI TÊN TÀI KHOẢN VÀ MẬT KHẨU", font=("Times New Roman", 18, 'bold'), bg="#e6f2ff", fg="#003366").pack(pady=20)
        tk.Label(change_frame, text="Vai trò:", font=("Times New Roman", 16, 'bold'), bg="#e6f2ff", fg="#FF0000").pack(pady=5)

        # Biến toàn cục cho role
        global role_var
        role_var = tk.StringVar(value="Teacher")

        # Tạo giao diện cho các nút radio
        radio_frame = tk.Frame(change_frame, bg="#e6f2ff")
        radio_frame.pack(pady=5)

        tk.Radiobutton(radio_frame, text="Giáo viên", variable=role_var, value="Teacher", font=("Times New Roman", 14, 'bold'), bg="#e6f2ff", fg="#0000FF").pack(side=tk.TOP, anchor=tk.CENTER)
        tk.Radiobutton(radio_frame, text="Học sinh", variable=role_var, value="Student", font=("Times New Roman", 14, 'bold'), bg="#e6f2ff", fg="#0000FF").pack(side=tk.TOP, anchor=tk.CENTER)

        # Tạo các Label và Entry cho thông tin tài khoản, Label được căn phải
        old_username_frame = tk.Frame(change_frame, bg="#e6f2ff")
        old_username_frame.pack(pady=5)
        global username_entry
        tk.Label(old_username_frame, text="Tên đăng nhập cũ:", font=("Times New Roman", 14, 'bold'), bg="#e6f2ff", fg="#003366", anchor='e', width=20).pack(side=tk.LEFT)
        username_entry = tk.Entry(old_username_frame, width=30, font=("Times New Roman", 12))
        username_entry.pack(side=tk.RIGHT, padx=10)

        old_password_frame = tk.Frame(change_frame, bg="#e6f2ff")
        old_password_frame.pack(pady=5)
        global password_entry
        tk.Label(old_password_frame, text="Mật khẩu cũ:", font=("Times New Roman", 14, 'bold'), bg="#e6f2ff", fg="#003366", anchor='e', width=20).pack(side=tk.LEFT)
        password_entry = tk.Entry(old_password_frame, width=30, font=("Times New Roman", 12), show="*")
        password_entry.pack(side=tk.RIGHT, padx=10)

        new_username_frame = tk.Frame(change_frame, bg="#e6f2ff")
        new_username_frame.pack(pady=5)
        global new_username_entry
        tk.Label(new_username_frame, text="Tên đăng nhập mới:", font=("Times New Roman", 14, 'bold'), bg="#e6f2ff", fg="#003366", anchor='e', width=20).pack(side=tk.LEFT)
        new_username_entry = tk.Entry(new_username_frame, width=30, font=("Times New Roman", 12))
        new_username_entry.pack(side=tk.RIGHT, padx=10)
        
        new_password_frame = tk.Frame(change_frame, bg="#e6f2ff")
        new_password_frame.pack(pady=5)
        global new_password_entry
        tk.Label(new_password_frame, text="Mật khẩu mới:", font=("Times New Roman", 14, 'bold'), bg="#e6f2ff", fg="#003366", anchor='e', width=20).pack(side=tk.LEFT)
        new_password_entry = tk.Entry(new_password_frame, width=30, font=("Times New Roman", 12), show="*")
        new_password_entry.pack(side=tk.RIGHT, padx=10)

        # Nút thay đổi tài khoản
        tk.Button(change_frame, text="Thay đổi tài khoản", command=handle_change, font=("Times New Roman", 14, 'bold'), bg="#003366", fg="white").pack(pady=20)
        tk.Button(change_frame, text="Thoát", command=change_frame.pack_forget, font=("Times New Roman", 14, 'bold'), bg="#800000", fg="white").pack(pady=10)
    else:
        change_frame.pack(fill="both", expand=True)  # Hiển thị lại change_frame nếu đã tồn tại


# NHẬP CÂU HỎI TRẮC NGHIỆM 4 LỰA CHỌN
# Hàm tạo file mẫu
def create_sample_file():
    doc = docx.Document()
    doc.add_heading('Mẫu câu hỏi trắc nghiệm nhiều lựa chọn', level=1)
    
    doc.add_paragraph('Câu hỏi:')
    doc.add_paragraph('Lựa chọn A:')
    doc.add_paragraph('Lựa chọn B:')
    doc.add_paragraph('Lựa chọn C:')
    doc.add_paragraph('Lựa chọn D:')
    doc.add_paragraph('Đáp án đúng:')
    doc.add_paragraph('Bài:')
    
    file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                           filetypes=[("Word files", "*.docx")],
                                           title="Lưu file mẫu")
    if file_path:
        doc.save(file_path)
        messagebox.showinfo("Thông báo", "File mẫu đã được tạo thành công!")

# Hàm xử lý khi chọn "Trắc nghiệm nhiều lựa chọn"
def save_question(question, option_a, option_b, option_c, option_d, correct_answer, option_level):
    # Kiểm tra trùng lặp đáp án
    options = [option_a, option_b, option_c, option_d]
    if len(options) != len(set(options)):
        messagebox.showerror("Lỗi", "Trùng đáp án")
        return

    # Kiểm tra quyền truy cập file
    if os.access("questions.xlsx", os.W_OK) or not os.path.exists("questions.xlsx"):
    # Mở hoặc tạo file Excel
        try:
            workbook = openpyxl.load_workbook("questions.xlsx")
        except FileNotFoundError:
            workbook = openpyxl.Workbook()
            workbook.remove(workbook.active)
            sheet = workbook.create_sheet("Questions")
            sheet.append(["Câu hỏi", "Lựa chọn A", "Lựa chọn B", "Lựa chọn C", "Lựa chọn D", "Đáp án đúng", "Bài"])
        else:
            sheet = workbook["Questions"]

        # Thêm câu hỏi vào file Excel
        sheet.append([question, option_a, option_b, option_c, option_d, correct_answer, option_level])
        workbook.save("questions.xlsx")
        return True  # Trả về True khi lưu thành công
    else:
        messagebox.showerror("Lỗi", "Không thể ghi vào file 'questions.xlsx'. Vui lòng kiểm tra quyền truy cập hoặc đóng file nếu nó đang được mở.")
# Hàm nhập câu hỏi từ file Word
def import_from_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if not file_path:
        return

    try:
        doc = docx.Document(file_path)
        questions = []
        current_question = {}

        for para in doc.paragraphs:
            if para.text.startswith("Câu hỏi:"):
                if current_question:
                    questions.append(current_question)
                    current_question = {}
                current_question['question'] = para.text.replace("Câu hỏi:", "").strip()
            elif para.text.startswith("Lựa chọn A:"):
                current_question['option_a'] = para.text.replace("Lựa chọn A:", "").strip()
            elif para.text.startswith("Lựa chọn B:"):
                current_question['option_b'] = para.text.replace("Lựa chọn B:", "").strip()
            elif para.text.startswith("Lựa chọn C:"):
                current_question['option_c'] = para.text.replace("Lựa chọn C:", "").strip()
            elif para.text.startswith("Lựa chọn D:"):
                current_question['option_d'] = para.text.replace("Lựa chọn D:", "").strip()
            elif para.text.startswith("Đáp án đúng:"):
                current_question['correct_answer'] = para.text.replace("Đáp án đúng:", "").strip()
            elif para.text.startswith("Bài:"):
                current_question['level'] = para.text.replace("Bài:", "").strip()

        # Thêm câu hỏi cuối cùng nếu có
        if current_question:
            questions.append(current_question)

        # Lưu tất cả câu hỏi vào file Excel
        for q in questions:
            save_question(q['question'], q['option_a'], q['option_b'], q['option_c'], q['option_d'], q['correct_answer'], q['level'])

        messagebox.showinfo("Thông báo", "Nhập câu hỏi từ file thành công!")
    except Exception as e:
        messagebox.showerror("Lỗi", f"Không thể nhập câu hỏi từ file: {e}")

def multiple_choice():
    exercise_frame = tk.Frame(root, bg="#e6f2ff")  # Khởi tạo test_paper_frame
    exercise_frame.pack(fill="both", expand=True)
    tk.Label(exercise_frame, text="TẠO CÂU HỎI TRẮC NGHIỆM 4 LỰA CHỌN", font=("Times_New_Roman", 18, "bold"), bg="#e6f2ff", fg="#003366").pack(pady=10)

    tk.Label(exercise_frame, text="Nội dung câu hỏi:", font=("Times_New_Roman", 14, "bold"), bg="#e6f2ff", fg="#003366").pack(anchor='w')
    question_text = tk.Text(exercise_frame, width=140, height=4, wrap=tk.WORD)
    question_text.pack(anchor='center')

    tk.Label(exercise_frame, text="Lựa chọn A:", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#003366").pack(anchor='w')
    option_a_text = tk.Text(exercise_frame, width=140, height=2, wrap=tk.WORD)
    option_a_text.pack(anchor='center')

    tk.Label(exercise_frame, text="Lựa chọn B:", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#003366").pack(anchor='w')
    option_b_text = tk.Text(exercise_frame, width=140, height=2, wrap=tk.WORD)
    option_b_text.pack(anchor='center')

    tk.Label(exercise_frame, text="Lựa chọn C:", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#003366").pack(anchor='w')
    option_c_text = tk.Text(exercise_frame, width=140, height=2, wrap=tk.WORD)
    option_c_text.pack(anchor='center')

    tk.Label(exercise_frame, text="Lựa chọn D:", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#003366").pack(anchor='w')
    option_d_text = tk.Text(exercise_frame, width=140, height=2, wrap=tk.WORD)
    option_d_text.pack(anchor='center')
    
    tk.Label(exercise_frame, text="Bài:", font=("Times_New_Roman", 14, "bold"), bg="#e6f2ff", fg="#003366").pack(anchor='w')
    option_level_text = tk.Text(exercise_frame, width=140, height=2, wrap=tk.WORD)
    option_level_text.pack(anchor='center')

    # Đặt thứ tự chuyển đổi khi nhấn phím Tab
    question_text.bind("<Tab>", lambda e: focus_next_widget(e, option_a_text))
    option_a_text.bind("<Tab>", lambda e: focus_next_widget(e, option_b_text))
    option_b_text.bind("<Tab>", lambda e: focus_next_widget(e, option_c_text))
    option_c_text.bind("<Tab>", lambda e: focus_next_widget(e, option_d_text))
    option_d_text.bind("<Tab>", lambda e: focus_next_widget(e, question_level_text))
    option_level_text.bind("<Tab>", lambda e: focus_next_widget(e, question_text))

    tk.Label(exercise_frame, text="Đáp án đúng:", font=("Times_New_Roman", 14, "bold"), bg="#e6f2ff", fg="#003366").pack(anchor='center')
    correct_answer_var = tk.StringVar(value="A")
    answer_frame = tk.Frame(exercise_frame)
    answer_frame.pack(anchor='center')
    tk.Radiobutton(answer_frame, text="A", variable=correct_answer_var, value="A", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#0000FF").pack(side='left')
    tk.Radiobutton(answer_frame, text="B", variable=correct_answer_var, value="B", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#0000FF").pack(side='left')
    tk.Radiobutton(answer_frame, text="C", variable=correct_answer_var, value="C", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#0000FF").pack(side='left')
    tk.Radiobutton(answer_frame, text="D", variable=correct_answer_var, value="D", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#0000FF").pack(side='left')

    button_frame = tk.Frame(exercise_frame)
    button_frame.pack(anchor='center', pady=10)
    tk.Button(button_frame, text="Lưu", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#FF0000", command=lambda: save_question(
        question_text.get("1.0", tk.END).strip(),
        option_a_text.get("1.0", tk.END).strip(),
        option_b_text.get("1.0", tk.END).strip(),
        option_c_text.get("1.0", tk.END).strip(),
        option_d_text.get("1.0", tk.END).strip(),
        correct_answer_var.get(),
        option_level_text.get("1.0", tk.END).strip()
    )).pack(side='left', padx=5)
    tk.Button(button_frame, text="Thoát", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#FF0000", command=exercise_frame.destroy).pack(side='left', padx=5)
    tk.Button(button_frame, text="Nhập từ File", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#FF0000", command=import_from_file).pack(side='left', padx=5)
    tk.Button(button_frame, text="Tải file mẫu", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#FF0000", command=create_sample_file).pack(side='left', padx=5)

def focus_next_widget(event, next_widget):
    event.widget.tk_focusNext().focus()
    return "break"

# NHẬP CÂU HỎI TRẮC NGHIỆM ĐÚNG/SAI
# Hàm tạo file mẫu
def create_sample_file_true_false():
    doc = docx.Document()
    doc.add_heading('Mẫu câu hỏi Đúng/Sai', level=1)
    
    doc.add_paragraph('Câu hỏi:')
    doc.add_paragraph('Lựa chọn A:')
    doc.add_paragraph('Lựa chọn B:')
    doc.add_paragraph('Lựa chọn C:')
    doc.add_paragraph('Lựa chọn D:')
    doc.add_paragraph('Đáp án A: Đúng/Sai')
    doc.add_paragraph('Đáp án B: Đúng/Sai')
    doc.add_paragraph('Đáp án C: Đúng/Sai')
    doc.add_paragraph('Đáp án D: Đúng/Sai')
    doc.add_paragraph('Bài: 1,2,3...')
    
    file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                             filetypes=[("Word files", "*.docx")],
                                             title="Lưu file mẫu")
    if file_path:
        doc.save(file_path)
        messagebox.showinfo("Thông báo", "File mẫu đã được tạo thành công!")

# Hàm lưu câu hỏi đúng/sai vào file Excel
def save_true_false_question(question, option_a, answer_a, option_b, answer_b, option_c, answer_c, option_d, answer_d, option_level):
    if not question or not option_a or not option_b or not option_c or not option_d or not option_level:
        messagebox.showerror("Lỗi", "Vui lòng nhập đầy đủ thông tin.")
        return
    
    if os.access("questions_true_false.xlsx", os.W_OK) or not os.path.exists("questions_true_false.xlsx"):
        try:
            workbook = openpyxl.load_workbook("questions_true_false.xlsx")
        except FileNotFoundError:
            workbook = openpyxl.Workbook()
            workbook.remove(workbook.active)
            sheet = workbook.create_sheet("Questions")
            sheet.append(["Câu hỏi", "Lựa chọn A", "Đáp án A", "Lựa chọn B", "Đáp án B", "Lựa chọn C", "Đáp án C", "Lựa chọn D", "Đáp án D", "Bài"])
        else:
            sheet = workbook["Questions"]

        sheet.append([question, option_a, answer_a, option_b, answer_b, option_c, answer_c, option_d, answer_d, option_level])
        workbook.save("questions_true_false.xlsx")
        return True  # Trả về True khi lưu thành công
    else:
        messagebox.showerror("Lỗi", "Không thể ghi vào file 'questions_true_false.xlsx'. Vui lòng kiểm tra quyền truy cập.")
        return False  # Trả về False khi có lỗi

# Hàm nhập câu hỏi từ file Word
def import_from_file_tf():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if not file_path:
        return

    try:
        doc = docx.Document(file_path)
        questions = []
        current_question = {}

        for para in doc.paragraphs:
            if para.text.startswith("Câu hỏi:"):
                if current_question:
                    questions.append(current_question)
                    current_question = {}
                current_question['question'] = para.text.replace("Câu hỏi:", "").strip()
            elif para.text.startswith("Lựa chọn A:"):
                current_question['option_a'] = para.text.replace("Lựa chọn A:", "").strip()
            elif para.text.startswith("Lựa chọn B:"):
                current_question['option_b'] = para.text.replace("Lựa chọn B:", "").strip()
            elif para.text.startswith("Lựa chọn C:"):
                current_question['option_c'] = para.text.replace("Lựa chọn C:", "").strip()
            elif para.text.startswith("Lựa chọn D:"):
                current_question['option_d'] = para.text.replace("Lựa chọn D:", "").strip()
            elif para.text.startswith("Đáp án A:"):
                current_question['answer_a'] = para.text.replace("Đáp án A:", "").strip()
            elif para.text.startswith("Đáp án B:"):
                current_question['answer_b'] = para.text.replace("Đáp án B:", "").strip()
            elif para.text.startswith("Đáp án C:"):
                current_question['answer_c'] = para.text.replace("Đáp án C:", "").strip()
            elif para.text.startswith("Đáp án D:"):
                current_question['answer_d'] = para.text.replace("Đáp án D:", "").strip()
            elif para.text.startswith("Bài:"):
                current_question['level'] = para.text.replace("Bài:", "").strip()

        if current_question:
            questions.append(current_question)

        for q in questions:
            save_true_false_question(q['question'], q['option_a'], q['answer_a'], q['option_b'], q['answer_b'], q['option_c'], q['answer_c'], q['option_d'], q['answer_d'], q['level'])

        messagebox.showinfo("Thông báo", "Nhập câu hỏi từ file thành công!")
    except Exception as e:
        messagebox.showerror("Lỗi", f"Không thể nhập câu hỏi từ file: {e}")

# Hàm tạo cửa sổ nhập câu hỏi đúng/sai
def true_false():
    def save():
        question = question_text.get("1.0", tk.END).strip()
        option_a = option_a_text.get("1.0", tk.END).strip()
        option_b = option_b_text.get("1.0", tk.END).strip()
        option_c = option_c_text.get("1.0", tk.END).strip()
        option_d = option_d_text.get("1.0", tk.END).strip()
        option_level = option_level_text.get("1.0", tk.END).strip()
        answer_a = answer_a_var.get()
        answer_b = answer_b_var.get()
        answer_c = answer_c_var.get()
        answer_d = answer_d_var.get()
        
        if save_true_false_question(question, option_a, answer_a, option_b, answer_b, option_c, answer_c, option_d, answer_d, option_level):
            messagebox.showinfo("Thông báo", "Câu hỏi đã được lưu thành công!")

    exercise_frame = tk.Frame(root, bg="#e6f2ff")  # Khởi tạo test_paper_frame
    exercise_frame.pack(fill="both", expand=True)

    tk.Label(exercise_frame, text="TẠO CÂU HỎI TRẮC NGHIỆM ĐÚNG/SAI", font=("Times_New_Roman", 18, "bold"), bg="#e6f2ff", fg="#003366").pack(pady=10)

    tk.Label(exercise_frame, text="Nội dung câu hỏi:", font=("Times_New_Roman", 14, "bold"), bg="#e6f2ff", fg="#003366").pack(anchor='w')
    question_text = tk.Text(exercise_frame, width=140, height=4, wrap=tk.WORD)
    question_text.pack(anchor='center')

    tk.Label(exercise_frame, text="Lựa chọn A:", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#003366").pack(anchor='w')
    option_a_text = tk.Text(exercise_frame, width=140, height=2, wrap=tk.WORD)
    option_a_text.pack(anchor='center')

    tk.Label(exercise_frame, text="Lựa chọn B:", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#003366").pack(anchor='w')
    option_b_text = tk.Text(exercise_frame, width=140, height=2, wrap=tk.WORD)
    option_b_text.pack(anchor='center')

    tk.Label(exercise_frame, text="Lựa chọn C:", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#003366").pack(anchor='w')
    option_c_text = tk.Text(exercise_frame, width=140, height=2, wrap=tk.WORD)
    option_c_text.pack(anchor='center')

    tk.Label(exercise_frame, text="Lựa chọn D:", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#003366").pack(anchor='w')
    option_d_text = tk.Text(exercise_frame, width=140, height=2, wrap=tk.WORD)
    option_d_text.pack(anchor='center')

    tk.Label(exercise_frame, text="Bài:", font=("Times_New_Roman", 14, "bold"), bg="#e6f2ff", fg="#003366").pack(anchor='w')
    option_level_text = tk.Text(exercise_frame, width=140, height=2, wrap=tk.WORD)
    option_level_text.pack(anchor='center')

   # Bố trí các nút chọn đáp án trên cùng một dòng
    answer_frame = tk.Frame(exercise_frame)
    answer_frame.pack(anchor='center', pady=5)

  # Đặt thứ tự chuyển đổi khi nhấn phím Tab
    question_text.bind("<Tab>", lambda e: focus_next_widget(e, option_a_text))
    option_a_text.bind("<Tab>", lambda e: focus_next_widget(e, option_b_text))
    option_b_text.bind("<Tab>", lambda e: focus_next_widget(e, option_c_text))
    option_c_text.bind("<Tab>", lambda e: focus_next_widget(e, option_d_text))
    option_d_text.bind("<Tab>", lambda e: focus_next_widget(e, option_level_text))
    option_level_text.bind("<Tab>", lambda e: focus_next_widget(e, question_text))
    
    tk.Label(answer_frame, text="Đáp án lựa chọn A:", font=("Times_New_Roman", 10, "bold"), bg="#e6f2ff", fg="#003366").pack(side='left')
    answer_a_var = tk.StringVar(value="Đúng")
    tk.Radiobutton(answer_frame, text="Đúng", font=("Times_New_Roman", 10, "bold"), bg="#e6f2ff", fg="#0000FF", variable=answer_a_var, value="Đúng").pack(side='left')
    tk.Radiobutton(answer_frame, text="Sai", font=("Times_New_Roman", 10, "bold"), bg="#e6f2ff", fg="#FF0000", variable=answer_a_var, value="Sai").pack(side='left')

    tk.Label(answer_frame, text="Đáp án lựa chọn B:", font=("Times_New_Roman", 10, "bold"), bg="#e6f2ff", fg="#003366").pack(side='left', padx=(10, 0))
    answer_b_var = tk.StringVar(value="Đúng")
    tk.Radiobutton(answer_frame, text="Đúng", font=("Times_New_Roman", 10, "bold"), bg="#e6f2ff", fg="#0000FF", variable=answer_b_var, value="Đúng").pack(side='left')
    tk.Radiobutton(answer_frame, text="Sai", font=("Times_New_Roman", 10, "bold"), bg="#e6f2ff", fg="#FF0000", variable=answer_b_var, value="Sai").pack(side='left')

    tk.Label(answer_frame, text="Đáp án lựa chọn C:", font=("Times_New_Roman", 10, "bold"), bg="#e6f2ff", fg="#003366").pack(side='left', padx=(10, 0))
    answer_c_var = tk.StringVar(value="Đúng")
    tk.Radiobutton(answer_frame, text="Đúng", font=("Times_New_Roman", 10, "bold"), bg="#e6f2ff", fg="#0000FF", variable=answer_c_var, value="Đúng").pack(side='left')
    tk.Radiobutton(answer_frame, text="Sai", font=("Times_New_Roman", 10, "bold"), bg="#e6f2ff", fg="#FF0000", variable=answer_c_var, value="Sai").pack(side='left')

    tk.Label(answer_frame, text="Đáp án lựa chọn D:", font=("Times_New_Roman", 10, "bold"), bg="#e6f2ff", fg="#003366").pack(side='left', padx=(10, 0))
    answer_d_var = tk.StringVar(value="Đúng")
    tk.Radiobutton(answer_frame, text="Đúng", font=("Times_New_Roman", 10, "bold"), bg="#e6f2ff", fg="#0000FF", variable=answer_d_var, value="Đúng").pack(side='left')
    tk.Radiobutton(answer_frame, text="Sai", font=("Times_New_Roman", 10, "bold"), bg="#e6f2ff", fg="#FF0000", variable=answer_d_var, value="Sai").pack(side='left')

    button_frame = tk.Frame(exercise_frame)
    button_frame.pack(anchor='center', pady=10)
    # Bố trí các nút lệnh
    tk.Button(button_frame, text="Lưu", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#FF0000", command=save).pack(side='left', padx=5)
    tk.Button(button_frame, text="Thoát", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#FF0000", command=exercise_frame.destroy).pack(side='left', padx=5)
    tk.Button(button_frame, text="Nhập từ file", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#FF0000", command=import_from_file_tf).pack(side='left', padx=5)
    tk.Button(button_frame, text="Tải file mẫu", font=("Times_New_Roman", 12, "bold"), bg="#e6f2ff", fg="#FF0000", command=create_sample_file_true_false).pack(side='left', padx=5)
    
# TRỘN ĐỀ TỪ NGÂN HÀNG CÂU HỎI
# Hàm tải câu hỏi trắc nghiệm 4 lựa chọn từ file Excel
def load_multiple_choice_questions(file_name, sheet_name):
    multiple_choice_questions = []
    if os.path.exists(file_name):
        try:
            workbook = openpyxl.load_workbook(file_name)
            sheet = workbook[sheet_name]
            for row in list(sheet.iter_rows(values_only=True))[1:]:
                if len(row) >= 6:
                    multiple_choice_questions.append({
                        'question': row[0],
                        'options': [str(row[1]).strip(), str(row[2]).strip(), str(row[3]).strip(), str(row[4]).strip()],
                        'correct_answer': str(row[5]).strip().lower(),
                    })
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file câu hỏi: {str(e)}")
            return []
    else:
        messagebox.showerror("Lỗi", f"File {file_name} không tồn tại.")
        return []
    return multiple_choice_questions

# Hàm tải câu hỏi đúng/sai từ file Excel
def load_true_false_questions(file_name, sheet_name):
    true_false_questions = []
    if os.path.exists(file_name):
        try:
            workbook = openpyxl.load_workbook(file_name)
            sheet = workbook[sheet_name]
            for row in list(sheet.iter_rows(values_only=True))[1:]:
                if len(row) >= 9:
                    true_false_questions.append({
                        'question': row[0],
                        'options': [str(row[1]).strip(), str(row[3]).strip(), str(row[5]).strip(), str(row[7]).strip()],
                        'correct_answers': [str(row[2]).strip().lower(), str(row[4]).strip().lower(), str(row[6]).strip().lower(), str(row[8]).strip().lower()],
                    })
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file câu hỏi: {str(e)}")
            return []
    else:
        messagebox.showerror("Lỗi", f"File {file_name} không tồn tại.")
        return []
    return true_false_questions

# Hàm lưu đề đã trộn vào file Word
def save_quiz_to_word(mc_questions, tf_questions, number_of_quizzes, path):
    # Lưu đề gốc
    doc = Document()
    doc.add_heading('Đề gốc', level=1)
    doc.add_heading('Phần I: Câu hỏi trắc nghiệm nhiều lựa chọn', level=2)
    for index, question in enumerate(mc_questions, start=1):
        doc.add_paragraph(f"Câu {index}: {question['question']}")
        options = question['options'][:]
        for idx, option in enumerate(options):
            doc.add_paragraph(f"{chr(65 + idx)}. {option}")  # A, B, C, D
        doc.add_paragraph(f"Câu trả lời đúng: {question['correct_answer']}")
        doc.add_paragraph()

    doc.add_heading('Phần II: Câu hỏi trắc nghiệm đúng/sai', level=2)
    for index, question in enumerate(tf_questions, start=1):
        doc.add_paragraph(f"Câu {index}: {question['question']}")
        options = question['options'][:]
        for idx, option in enumerate(options):
            doc.add_paragraph(f"{chr(65 + idx)}. {option}")  # A, B, C, D
        doc.add_paragraph(f"Câu trả lời đúng: {', '.join(question['correct_answers'])}")
        doc.add_paragraph()

    doc.save(os.path.join(path, "De_goc.docx"))

    # Lưu các đề con
    for i in range(number_of_quizzes):
        doc = Document()
        doc.add_heading(f'Đề thi {i + 1}', level=1)

        # Phần I: Câu hỏi trắc nghiệm nhiều lựa chọn
        doc.add_heading('Phần I: Câu hỏi trắc nghiệm nhiều lựa chọn', level=2)
        shuffled_mc = random.sample(mc_questions, len(mc_questions))
        mc_question_number = 1  # Bắt đầu từ Câu 1 cho mỗi đề
        for question in shuffled_mc:
            doc.add_paragraph(f"Câu {mc_question_number}: {question['question']}")
            options = question['options'][:]
            random.shuffle(options)  # Trộn các đáp án
            for idx, option in enumerate(options):
                doc.add_paragraph(f"{chr(65 + idx)}. {option}")  # A, B, C, D
            
            # Ghi câu trả lời đúng dựa trên đáp án đã trộn
            correct_answer_index = options.index(question['correct_answer']) if question['correct_answer'] in options else -1
            if correct_answer_index != -1:
                doc.add_paragraph(f"Câu trả lời đúng: {chr(65 + correct_answer_index)}")  # A, B, C, D
            doc.add_paragraph()  # Thêm dòng trống
            mc_question_number += 1

        # Phần II: Câu hỏi đúng/sai
        doc.add_heading('Phần II: Câu hỏi trắc nghiệm đúng/sai', level=2)
        shuffled_tf = random.sample(tf_questions, len(tf_questions))
        tf_question_number = 1  # Bắt đầu từ Câu 1 cho mỗi đề
        for question in shuffled_tf:
            doc.add_paragraph(f"Câu {tf_question_number}: {question['question']}")
            options = question['options'][:]
            random.shuffle(options)  # Trộn các đáp án
            for idx, option in enumerate(options):
                doc.add_paragraph(f"{chr(65 + idx)}. {option}")  # A, B, C, D

            # Ghi câu trả lời đúng dựa trên đáp án đã trộn
            correct_answers_indices = [options.index(ans) for ans in question['correct_answers'] if ans in options]
            if correct_answers_indices:
                correct_answers_labels = [chr(65 + index) for index in correct_answers_indices]
                doc.add_paragraph(f"Câu trả lời đúng: {', '.join(correct_answers_labels)}")  # A, B, C, D
            doc.add_paragraph()  # Thêm dòng trống
            tf_question_number += 1

        doc.save(os.path.join(path, f"De_{i + 1:03d}.docx"))  # Tạo tên file De_001.docx, De_002.docx,...

# Hàm trộn đề thi
def shuffle_quiz():
    global number_of_quizzes, mc_questions, tf_questions, save_path
    try:
        number_of_quizzes = int(num_quizzes_entry.get().strip())
        if number_of_quizzes <= 0:
            raise ValueError("Số đề phải lớn hơn 0.")
    except ValueError as e:
        messagebox.showerror("Lỗi", f"Vui lòng nhập số nguyên dương cho số đề: {e}")
        return

    mc_questions = load_multiple_choice_questions("questions.xlsx", "Questions")
    tf_questions = load_true_false_questions("questions_true_false.xlsx", "Questions")

    if not mc_questions and not tf_questions:
        messagebox.showerror("Lỗi", "Chưa có câu hỏi nào được nhập hoặc định dạng không đúng.")
        return

    quiz_text = ""
    # Ghi đề gốc
    quiz_text += "Đề gốc:\n\n"

    # Phần I: Câu hỏi trắc nghiệm nhiều lựa chọn
    quiz_text += "Phần I: Câu hỏi trắc nghiệm nhiều lựa chọn\n"
    for index, question in enumerate(mc_questions, start=1):
        quiz_text += f"Câu {index}: {question['question']}\n"
        options = question['options'][:]
        for idx, option in enumerate(options):
            quiz_text += f"{chr(65 + idx)}. {option}\n"  # A, B, C, D
        quiz_text += f"Câu trả lời đúng: {question['correct_answer']}\n\n"

    # Phần II: Câu hỏi đúng/sai
    quiz_text += "Phần II: Câu hỏi trắc nghiệm đúng/sai\n"
    for index, question in enumerate(tf_questions, start=1):
        quiz_text += f"Câu {index}: {question['question']}\n"
        options = question['options'][:]
        for idx, option in enumerate(options):
            quiz_text += f"{chr(65 + idx)}. {option}\n"  # A, B, C, D
        quiz_text += f"Câu trả lời đúng: {', '.join(question['correct_answers'])}\n\n"

    # Lưu các đề con vào file Word
    if save_path:
        save_quiz_to_word(mc_questions, tf_questions, number_of_quizzes, save_path)
        messagebox.showinfo("Thông báo", "Các đề đã được lưu thành công!")
    else:
        messagebox.showerror("Lỗi", "Vui lòng chọn đường dẫn lưu.")

# Hàm chọn đường dẫn lưu
def choose_save_path():
    global save_path
    save_path = filedialog.askdirectory()
    if save_path:
        path_label.config(text=f"Đường dẫn lưu: {save_path}")

# Hàm mở cửa sổ trộn đề thi
def open_shuffle_window():
    exercise_frame = tk.Frame(root, bg="#e6f2ff")  # Khởi tạo test_paper_frame
    exercise_frame.pack(fill="both", expand=True)
    def on_close():
        exercise_frame.grab_release()
        exercise_frame.destroy()
    # Nhập số đề
    num_quizzes_label = tk.Label(exercise_frame, text="TRỘN ĐỀ TỪ NGÂN HÀNG CÂU HỎI", font=("Times New Roman", 18, "bold"), bg="#e6f2ff", fg="#003366")
    num_quizzes_label.pack(pady=10)
    num_quizzes_label = tk.Label(exercise_frame, text="Nhập số đề cần trộn:", font=("Times New Roman", 14, "bold"), bg="#e6f2ff", fg="#0000FF")
    num_quizzes_label.pack(pady=10)
    global num_quizzes_entry
    num_quizzes_entry = tk.Entry(exercise_frame, font=("Times New Roman", 14))
    num_quizzes_entry.pack(pady=5)

    # Nút chọn đường dẫn lưu
    global path_label
    path_label = tk.Label(exercise_frame, text="Bạn chưa chọn đường dẫn lưu.", font=("Times New Roman", 12, "bold"), bg="#e6f2ff", fg="#FF0000")
    path_label.pack(pady=10)
    choose_path_button = tk.Button(exercise_frame, text="Chọn đường dẫn lưu", font=("Times New Roman", 14, "bold"), fg="#0000FF", command=choose_save_path)
    choose_path_button.pack(pady=5)

    # Nút trộn đề
    shuffle_button = tk.Button(exercise_frame, text="Trộn Đề", font=("Times New Roman", 14, "bold"), bg="#32CD32", fg="#FFFFFF", command=shuffle_quiz)
    shuffle_button.pack(pady=20)

    # Nút thoát
    exit_button = tk.Button(exercise_frame, text="Thoát", font=("Times New Roman", 14, "bold"), bg="#FF0000", fg="#FFFFFF", command=on_close)
    exit_button.pack(pady=10)
    
# TRỘN ĐỀ TỪ FILE WORD
class Question:
    def __init__(self, question_text, options=None, correct_answers=None, question_type="multiple_choice"):
        self.question_text = question_text
        self.options = options or []
        self.correct_answers = correct_answers or []
        self.question_type = question_type

    def shuffle_options(self):
        original_correct_answers = self.correct_answers.copy()  # Lưu danh sách đáp án đúng gốc
        random.shuffle(self.options)  # Trộn đáp án
        # Cập nhật lại đáp án đúng dựa trên vị trí mới của đáp án đã trộn
        self.correct_answers = [self.options.index(self.options[i]) for i in original_correct_answers]

    def format_options(self):
        if self.question_type == "multiple_choice":
            labels = ['A', 'B', 'C', 'D']
            return [f"{labels[i]}. {option}" for i, option in enumerate(self.options)]
        elif self.question_type == "true_false":
            labels = ['a)', 'b)', 'c)', 'd)']
            return [f"{labels[i]} {option}" for i, option in enumerate(self.options)]
        else:
            return self.options

class Quiz:
    def __init__(self):
        self.multiple_choice_questions = []
        self.true_false_questions = []

    def add_multiple_choice_question(self, question_text, options, correct_answer):
        if len(options) != 4:
            raise ValueError("Options must be a list of 4 choices.")
        question = Question(question_text, options, [int(correct_answer) - 1], question_type="multiple_choice")
        self.multiple_choice_questions.append(question)

    def add_true_false_question(self, question_text, options, correct_answers):
        if len(options) != 4:
            raise ValueError("Options must be a list of 4 choices.")
        question = Question(question_text, options, correct_answers, question_type="true_false")
        self.true_false_questions.append(question)

    def shuffle_questions(self):
        random.shuffle(self.multiple_choice_questions)
        random.shuffle(self.true_false_questions)

    def display_quiz(self):
        output = []

        # Phần I: Trắc nghiệm 4 lựa chọn
        output.append("Phần I. Trắc nghiệm 4 lựa chọn:")
        for idx, question in enumerate(self.multiple_choice_questions):
            question.shuffle_options()  # Trộn đáp án
            question_display = f"Câu {idx + 1}: {question.question_text}\n" + "\n".join(question.format_options()) + "\n"
            output.append(question_display)

        # Phần II: Trắc nghiệm đúng/sai
        output.append("Phần II. Trắc nghiệm đúng/sai:")
        for idx, question in enumerate(self.true_false_questions):
            question.shuffle_options()  # Trộn đáp án
            question_display = f"Câu {idx + 1}: {question.question_text}\n" + "\n".join(question.format_options()) + "\n"
            output.append(question_display)

        return output

def create_sample_file_w():
    # Chọn nơi lưu file mẫu
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")], title="Chọn nơi lưu file mẫu")
    if file_path:
        # Tạo file mẫu
        doc = Document()
        doc.add_heading('Hướng dẫn nhập câu hỏi', level=1)

        # Ví dụ câu hỏi trắc nghiệm nhiều lựa chọn
        doc.add_paragraph("Hành tinh nào gần Mặt Trời nhất?")
        doc.add_paragraph("Sao Thủy")
        doc.add_paragraph("Sao Kim")
        doc.add_paragraph("Sao Hỏa")
        doc.add_paragraph("Trái Đất")
        doc.add_paragraph("1")  # Đáp án đúng

        # Ví dụ câu hỏi trắc nghiệm đúng/sai
        doc.add_paragraph("Dưới đây là các hành tinh trong hệ mặt trời:")
        doc.add_paragraph("a) Trái Đất")
        doc.add_paragraph("b) Mặt Trời")
        doc.add_paragraph("c) Sao Hỏa")
        doc.add_paragraph("d) Sao Kim")
        doc.add_paragraph("0, 2")  # Đáp án đúng: a) và c)

        # Lưu file mẫu
        doc.save(file_path)
        messagebox.showinfo("Thông báo", f"File mẫu đã được lưu tại: {file_path}")

def load_questions_from_docx(file_path):
    questions = []
    doc = Document(file_path)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:  # Bỏ qua đoạn trống
            questions.append(text)
    return questions

def import_questions():
    file_path = filedialog.askopenfilename(title="Chọn file Word", filetypes=[("Word files", "*.docx")])
    if file_path:
        try:
            questions = load_questions_from_docx(file_path)
            i = 0
            while i < len(questions):
                question_text = questions[i]  # Câu hỏi
                options = questions[i + 1:i + 5]  # 4 đáp án
                correct_answer = questions[i + 5].strip()  # Đáp án đúng

                if correct_answer.isdigit() and int(correct_answer) in [1, 2, 3, 4]:  # Trắc nghiệm nhiều lựa chọn
                    quiz.add_multiple_choice_question(question_text, options, correct_answer)
                    i += 6  # Chuyển đến câu hỏi tiếp theo
                elif correct_answer:  # Kiểm tra dạng đúng/sai
                    correct_answers = list(map(int, correct_answer.split(',')))  # Lấy danh sách đáp án đúng
                    quiz.add_true_false_question(question_text, options, correct_answers)
                    i += 6  # Chuyển đến câu hỏi tiếp theo
                else:
                    messagebox.showwarning("Cảnh báo", f"Đáp án đúng '{correct_answer}' không hợp lệ.")
                    i += 6  # Bỏ qua câu hỏi không hợp lệ

            messagebox.showinfo("Thông báo", "Đã nhập câu hỏi thành công!")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể nhập câu hỏi từ file: {e}")

def save_quiz_to_word_w(quiz_output, file_path):
    doc = Document()
    doc.add_paragraph(quiz_output)
    doc.save(file_path)

def shuffle_and_display_quiz():
    try:
        num_tests = int(num_tests_entry.get().strip())
        if num_tests <= 0:
            raise ValueError("Số đề trộn phải lớn hơn 0.")
        
        # Trộn và lưu các đề
        for i in range(num_tests):
            shuffled_output = f"Đề số {i + 1}:\n"
            shuffled_output += "Phần I. Trắc nghiệm 4 lựa chọn:\n"
            random.shuffle(quiz.multiple_choice_questions)  # Trộn câu hỏi phần I
            for idx, question in enumerate(quiz.multiple_choice_questions):
                question.shuffle_options()  # Trộn đáp án mỗi câu hỏi
                question_display = f"Câu {idx + 1}: {question.question_text}\n" + "\n".join(question.format_options()) + "\n"
                shuffled_output += question_display
            
            shuffled_output += "Phần II. Trắc nghiệm đúng/sai:\n"
            random.shuffle(quiz.true_false_questions)  # Trộn câu hỏi phần II
            for idx, question in enumerate(quiz.true_false_questions):
                question.shuffle_options()  # Trộn đáp án mỗi câu hỏi
                question_display = f"Câu {idx + 1}: {question.question_text}\n" + "\n".join(question.format_options()) + "\n"
                shuffled_output += question_display
            
            # Lưu đề vào file riêng
            test_file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")],
                                                            title=f"Lưu Đề {i + 1}", initialfile=f"Đề {i + 1}.docx")
            if test_file_path:
                save_quiz_to_word_w(shuffled_output, test_file_path)
                messagebox.showinfo("Thông báo", f"Đề {i + 1} đã được lưu tại: {test_file_path}")

    except ValueError as e:
        messagebox.showerror("Lỗi", str(e))
def open_shuffle_window_word():
    exercise_frame = tk.Frame(root, bg="#e6f2ff")  # Khởi tạo test_paper_frame
    exercise_frame.pack(fill="both", expand=True)
    def on_close():
        exercise_frame.grab_release()
        exercise_frame.destroy()
    # Tải câu hỏi từ file Word
    num_tests_label = tk.Label(exercise_frame, text="TRỘN ĐỀ TỪ FILE CÂU HỎI", font=("Times New Roman", 18, "bold"), bg="#e6f2ff", fg="#003366")
    num_tests_label.pack(pady=10)
    load_button = tk.Button(exercise_frame, text="Tải câu hỏi từ file Word", command=import_questions, font=("Times New Roman", 14, "bold"), bg="#4CAF50", fg="white", width=20, height=2)
    load_button.pack(pady=10)
  
    # Nhập số đề trộn
    num_tests_label = tk.Label(exercise_frame, text="Nhập số đề trộn:", font=("Times New Roman", 14, "bold"), bg="#e6f2ff", fg="#0000FF")
    num_tests_label.pack(pady=5)
    global num_tests_entry
    num_tests_entry = tk.Entry(exercise_frame, font=("Times New Roman", 14), width=10)
    num_tests_entry.pack(pady=5)

    # Trộn và hiển thị đề
    shuffle_button = tk.Button(exercise_frame, text="Trộn đề", command=shuffle_and_display_quiz, font=("Times New Roman", 14, "bold"), bg="#FF5722", fg="white", width=20, height=2)
    shuffle_button.pack(pady=10)
    # Tạo file mẫu
    sample_button = tk.Button(exercise_frame, text="Tạo file mẫu câu hỏi", command=create_sample_file_w, font=("Times New Roman", 14, "bold"), bg="#008CBA", fg="white")
    sample_button.pack(pady=10)
    # Nút thoát
    exit_button = tk.Button(exercise_frame, text="Thoát", font=("Times New Roman", 14, "bold"), bg="#FF0000", fg="#FFFFFF", command=on_close)
    exit_button.pack(pady=10)
    
# BÀI DẠY TRẮC NGHIỆM
# Hàm bắt đầu bài tập trắc nghiệm
def start_quiz_bd(selected_unit):
    questions = []
    answers_vars = []
    score_label = None  # Khởi tạo biến để lưu label điểm số
    option_buttons = [[] for _ in range(4)]  # Danh sách lưu các nút lựa chọn

    # Hàm tải câu hỏi từ file Excel
    def load_multiple_choice_questions(file_name, sheet_name):
        if os.path.exists(file_name):
            try:
                workbook = openpyxl.load_workbook(file_name)
                sheet = workbook[sheet_name]
                for row in list(sheet.iter_rows(values_only=True))[1:]:  # Bỏ qua dòng tiêu đề
                    if len(row) >= 7:  # Kiểm tra nếu hàng có đủ 7 cột
                        question = row[0]
                        options = [
                            str(row[1]).strip(),  # Lựa chọn A
                            str(row[2]).strip(),  # Lựa chọn B
                            str(row[3]).strip(),  # Lựa chọn C
                            str(row[4]).strip()   # Lựa chọn D
                        ]
                        correct_answer = str(row[5]).strip().lower()  # Đáp án đúng
                        unit = str(row[6]).strip()  # Bài
                        
                        if f"Bài {unit}" == selected_unit:  # Chỉ thêm câu hỏi thuộc đơn vị được chọn
                            questions.append({
                                'question': question,
                                'options': options,
                                'correct_answer': correct_answer,
                            })
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể đọc file câu hỏi: {str(e)}")
                return False
        else:
            messagebox.showerror("Lỗi", f"File {file_name} không tồn tại.")
            return False
        return True

    # Tải câu hỏi từ file `questions.xlsx`
    if not load_multiple_choice_questions("questions.xlsx", "Questions"):
        return

    if not questions:
        messagebox.showerror("Lỗi", "Chưa có câu hỏi nào được nhập hoặc định dạng không đúng.")
        return

    # Tạo cửa sổ bài kiểm tra
    quiz_window = tk.Toplevel(root)
    quiz_window.title("Bài dạy trắc nghiệm")
    quiz_window.geometry("800x550")  # Thay đổi kích thước cửa sổ

    quiz_window.transient(root)  # Cửa sổ mới là cửa sổ con của root
    quiz_window.grab_set()  # Đảm bảo chỉ có thể thao tác với cửa sổ này

    # Căn giữa cửa sổ
    center_window(800, 550, quiz_window)
    # Trật tự cửa sổ
    window_order(quiz_window)

    # Tạo canvas cho phép thêm thanh cuộn
    canvas = tk.Canvas(quiz_window)
    canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    scrollbar = ttk.Scrollbar(quiz_window, orient="vertical", command=canvas.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    canvas.configure(yscrollcommand=scrollbar.set)
    canvas.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

    frame = tk.Frame(canvas)
    canvas.create_window((0, 0), window=frame, anchor="nw")

    # Căn giữa tiêu đề
    tk.Label(frame, text=f"BÀI DẠY TRẮC NGHIỆM: {selected_unit}", font=("Times New Roman", 18, "bold"), bg="#F0F0F0", fg="#003366").pack(pady=10)

    # Label để hiển thị số điểm
    score_label = tk.Label(frame, text="", font=("Times New Roman", 14, "bold"), fg="blue", bg="#F0F0F0")
    score_label.pack(anchor='nw', padx=20, pady=10)

    # Hàm kiểm tra và thay đổi màu sắc khi chọn đáp án
    def check_answer(question_idx, user_answer):
        correct_answer = questions[question_idx]['correct_answer']  # Đáp án đúng

        # So sánh đáp án của người dùng với đáp án đúng
        for idx, btn in enumerate(option_buttons):
            option_value = chr(97 + idx)  # Tính ra a, b, c, d
            if option_value == correct_answer:
                option_buttons[idx][question_idx].config(fg="green")  # Đáp án đúng -> xanh
            if option_value == user_answer:
                if user_answer != correct_answer:
                    option_buttons[idx][question_idx].config(fg="red")  # Đáp án sai -> đỏ

    # Hiển thị từng câu hỏi và các lựa chọn
    for i, question in enumerate(questions):
        question_text = question['question']
        options = question['options']

        tk.Label(frame, text=f"Câu {i + 1}: {question_text}", font=("Times New Roman", 12), bg="#F0F0F0").pack(anchor='w', padx=20, pady=5)

        # Biến lưu đáp án của người dùng
        answer_var = tk.StringVar(value="")
        answers_vars.append(answer_var)

        # Hiển thị các lựa chọn dạng RadioButton
        for idx, option in enumerate(options):
            rb = tk.Radiobutton(frame, text=f"{chr(97 + idx)}) {option}", variable=answer_var, value=chr(97 + idx),
                                font=("Times New Roman", 12), bg="#F0F0F0",
                                command=lambda idx=idx, i=i: check_answer(i, chr(97 + idx)))
            rb.pack(anchor='w', padx=40)
            option_buttons[idx].append(rb)  # Thêm nút vào danh sách

    # Hàm xử lý khi nộp bài
    def submit_answers(answers_vars):
        correct_count = 0
        total_questions = len(questions)

        for i, question in enumerate(questions):
            correct_answer = question['correct_answer']
            user_answer = answers_vars[i].get()

            if user_answer == correct_answer:
                correct_count += 1

        # Tính tỷ lệ đúng
        percentage = (correct_count / total_questions) * 10

        # Hiển thị kết quả
        score_label.config(text=f"Số câu đúng: {correct_count}/{total_questions}     Điểm: {round(percentage, 1)}", font=("Times New Roman", 18, "bold"))
      
    # Xem số câu đúng và điểm
    tk.Button(quiz_window, text="Xem số câu đúng và điểm", command=lambda: submit_answers(answers_vars), font=("Times New Roman", 16, "bold"), bg="#4CAF50", fg="white", activebackground="#45a049").pack(pady=20, anchor='center')

    # Nút thoát
    tk.Button(quiz_window, text="Thoát", command=quiz_window.destroy, font=("Times New Roman", 16, "bold"), bg="#800000", fg="white").pack(pady=10, anchor='center')

# Hàm tải danh sách đơn vị bài từ file
def load_units_bd(file_name, sheet_name):
    if os.path.exists(file_name):
        try:
            workbook = openpyxl.load_workbook(file_name)
            sheet = workbook[sheet_name]
            units = set()  # Sử dụng set để loại bỏ trùng lặp
            for row in list(sheet.iter_rows(values_only=True))[1:]:  # Bỏ qua dòng tiêu đề
                if len(row) >= 7:  # Kiểm tra nếu hàng có đủ 7 cột
                    unit = str(row[6]).strip()
                    units.add(unit)
            return sorted([f"Bài {unit}" for unit in units])  # Sắp xếp và định dạng cho combobox
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file đơn vị: {str(e)}")
            return []
    else:
        messagebox.showerror("Lỗi", f"File {file_name} không tồn tại.")
        return []

# Hàm khởi tạo cửa sổ chọn đơn vị bài
def teacher_menu_quiz_lesson():
    exercise_frame = tk.Frame(root, bg="#e6f2ff")  # Khởi tạo test_paper_frame
    exercise_frame.pack(fill="both", expand=True)
    tk.Label(exercise_frame, text="BÀI DẠY TRẮC NGHIỆM", font=("Times New Roman", 18, "bold"), bg="#e6f2ff", fg="#003366").pack(pady=10)
    # Tạo combobox chọn đơn vị bài
    tk.Label(exercise_frame, text="Chọn đơn vị bài:", font=("Times New Roman", 16, "bold"), bg="#e6f2ff", fg="#FF0000").pack(pady=10)
    global unit_var
    unit_var = tk.StringVar()
    unit_combobox = ttk.Combobox(exercise_frame, textvariable=unit_var, font=("Times New Roman", 14, "bold"))

    # Tải danh sách đơn vị từ file
    units = load_units_bd("questions.xlsx", "Questions")
    unit_combobox['values'] = units
    unit_combobox.pack(pady=10)
    unit_combobox.current(0)  # Chọn đơn vị đầu tiên làm mặc định

    # Nút bắt đầu bài kiểm tra
    start_button = tk.Button(exercise_frame, text="Bắt đầu làm bài", command=lambda: start_quiz_bd(unit_var.get()), font=("Times New Roman", 14, "bold"), bg="#4CAF50", fg="white", activebackground="#45a049")
    start_button.pack(pady=20)

    # Nút thoát
    tk.Button(exercise_frame, text="Thoát", command=exercise_frame.destroy, font=("Times New Roman", 14, "bold"), bg="#800000", fg="white").pack(pady=10)


# HIỂN THỊ KẾT QUẢ TRONG FILE EXCELL
#Hàm hiển thị kết quả trong file Excel
def show_results():
    file_name = "results.xlsx"
    if os.path.exists(file_name):
        workbook = openpyxl.load_workbook(file_name)
        sheet = workbook["Results"]

        result_window = tk.Toplevel(root)
        result_window.title("Kết quả bài kiểm tra")
        result_window.geometry("950x500")
        result_window.configure(bg="#f0f0f0")
        # Căn giữa cửa sổ
        center_window(950, 500, result_window)
        tk.Label(result_window, text="DANH SÁCH KẾT QUẢ", font=("Times New Roman", 18, "bold"), fg="#0000FF").pack(pady=10)
        # Hiển thị bảng kết quả
        result_frame = tk.Frame(result_window, bg="#f0f0f0")
        result_frame.pack(fill=tk.BOTH, expand=True)

        # Tạo bảng hiển thị kết quả
        columns = ["", "", "", "", ""]
        tree = ttk.Treeview(result_frame, columns=columns, show="headings", height=10)
        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, anchor='center')

        # Thêm dữ liệu vào bảng
        for row in sheet.iter_rows(values_only=True):
            tree.insert("", "end", values=row)

        tree.pack(fill=tk.BOTH, expand=True)
    else:
        messagebox.showerror("Lỗi", "Chưa có kết quả nào được lưu!")
    # Nút thoát
    tk.Button(result_window, text="Thoát", command=result_window.destroy, font=("Times New Roman", 16, "bold"), bg="#800000", fg="white").pack(pady=10, anchor='center')

# TẢI BẢNG KẾT QUẢ
# Hàm tải bảng kết quả về máy dưới dạng file Excel
def download_results():
    file_name = "results.xlsx"
    if os.path.exists(file_name):
        # Hiển thị hộp thoại để người dùng chọn nơi lưu file
        save_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if save_path:
            try:
                shutil.copy(file_name, save_path)
                messagebox.showinfo("Thành công", "File kết quả đã được tải về thành công!")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể tải file: {str(e)}")
    else:
        messagebox.showerror("Lỗi", "Chưa có kết quả nào để tải về!")

# BÀI TẬP TRẮC NGHIỆM 4 LỰA CHON
# Hàm kiểm tra thông tin nhập vào (Họ tên, Lớp) trước khi làm bài
def check_student_info():
    name = name_entry.get().strip()
    student_class = class_entry.get().strip()
    selected_unit = unit_var.get()

    if not name or not student_class:
        messagebox.showerror("Lỗi", "Vui lòng nhập đầy đủ Họ tên và Lớp!")
    else:
        start_quiz(name, student_class, selected_unit)

# Hàm xử lý bắt đầu bài tập trắc nghiệm
def start_quiz(name, student_class, selected_unit):
    questions = []
    answers_vars = []
    score_label = None  # Khởi tạo biến để lưu label điểm số

    # Hàm tải câu hỏi từ file Excel
    def load_multiple_choice_questions(file_name, sheet_name):
        if os.path.exists(file_name):
            try:
                workbook = openpyxl.load_workbook(file_name)
                sheet = workbook[sheet_name]
                for row in list(sheet.iter_rows(values_only=True))[1:]:  # Bỏ qua dòng tiêu đề
                    if len(row) >= 7:  # Kiểm tra nếu hàng có đủ 7 cột
                        question = row[0]
                        options = [
                            str(row[1]).strip(),  # Lựa chọn A
                            str(row[2]).strip(),  # Lựa chọn B
                            str(row[3]).strip(),  # Lựa chọn C
                            str(row[4]).strip()   # Lựa chọn D
                        ]
                        correct_answer = str(row[5]).strip().lower()  # Đáp án đúng
                        unit = str(row[6]).strip()  # Bài
                        
                        if f"Bài {unit}" == selected_unit:  # Chỉ thêm câu hỏi thuộc đơn vị được chọn
                            questions.append({
                                'question': question,
                                'options': options,
                                'correct_answer': correct_answer,
                            })
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể đọc file câu hỏi: {str(e)}")
                return False
        else:
            messagebox.showerror("Lỗi", f"File {file_name} không tồn tại.")
            return False
        return True

    # Tải câu hỏi từ file `questions.xlsx`
    if not load_multiple_choice_questions("questions.xlsx", "Questions"):
        return

    if not questions:
        messagebox.showerror("Lỗi", "Chưa có câu hỏi nào được nhập hoặc định dạng không đúng.")
        return

    # Tạo cửa sổ bài kiểm tra
    quiz_window = tk.Toplevel(root)
    quiz_window.title("Bài kiểm tra trắc nghiệm")
    quiz_window.geometry("1000x550")  # Thay đổi kích thước cửa sổ
    quiz_window.transient(root)  # Cửa sổ mới là cửa sổ con của root
    quiz_window.grab_set()  # Đảm bảo chỉ có thể thao tác với cửa sổ này
    # Căn giữa cửa sổ
    center_window(1000, 550, quiz_window)
    # Trật tự cửa sổ
    window_order(quiz_window)
    # Tạo canvas cho phép thêm thanh cuộn
    canvas = tk.Canvas(quiz_window)
    canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    scrollbar = ttk.Scrollbar(quiz_window, orient="vertical", command=canvas.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    canvas.configure(yscrollcommand=scrollbar.set)
    canvas.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

    frame = tk.Frame(canvas)
    canvas.create_window((0, 0), window=frame, anchor="nw")

    tk.Label(frame, text="THÔNG TIN NGƯỜI HỌC", font=("Times New Roman", 18, "bold"), fg="#003366").pack(pady=10)
    # Căn giữa tiêu đề và thông tin học sinh
    tk.Label(frame, text=f"Họ tên: {name}", font=("Times_New_Roman", 14, "bold"), bg="#F0F0F0", fg="#003366").pack(pady=5)
    tk.Label(frame, text=f"Lớp: {student_class}", font=("Times_New_Roman", 14, "bold"), bg="#F0F0F0", fg="#003366").pack(pady=5)
    tk.Label(frame, text=f"Bài tập trắc nghiệm: {selected_unit}", font=("Times_New_Roman", 16, "bold"), bg="#F0F0F0", fg="#0000FF").pack(pady=10)

    # Label để hiển thị số điểm
    score_label = tk.Label(frame, text="", font=("Times_New_Roman", 14, "bold"), fg="blue", bg="#F0F0F0")
    score_label.pack(anchor='center', padx=20, pady=10)

    def submit_answers():
        correct_count = 0
        total_questions = len(questions)
        unanswered_questions = []  # Danh sách câu chưa trả lời

        for i, question in enumerate(questions):
            correct_answer = question['correct_answer']  # Đáp án đúng
            user_answer = answers_vars[i].get().lower()  # Đáp án của người dùng

            # Kiểm tra nếu người dùng đã chọn đáp án
            if user_answer == "":
                unanswered_questions.append(i + 1)  # Lưu câu chưa trả lời theo thứ tự (1, 2, 3,...)
                continue  # Bỏ qua vòng lặp này và tiếp tục

            # So sánh đáp án của người dùng với đáp án đúng
            if user_answer == correct_answer:
                correct_count += 1
            # Tính tỷ lệ đúng
            percentage = (correct_count / total_questions) * 10

            # Thay đổi màu sắc của các lựa chọn
            for option in range(4):
                option_button[option][i].config(fg="black")  # Đặt màu mặc định cho tất cả
                if chr(97 + option) == correct_answer:  # Nếu là đáp án đúng
                    option_button[option][i].config(fg="green")
                elif user_answer == chr(97 + option):  # Nếu là đáp án sai
                    option_button[option][i].config(fg="red")

        # Kiểm tra nếu có câu chưa trả lời
        if unanswered_questions:
            questions_list = ", ".join(map(str, unanswered_questions))
            messagebox.showwarning("Cảnh báo", f"Các câu {questions_list} chưa được chọn đáp án. Vui lòng chọn đầy đủ đáp án.")
            return  # Dừng thực hiện hàm nếu còn câu chưa trả lời
        # Hiển thị kết quả
        score_label.config(text=f"Số câu đúng: {correct_count}/{total_questions}     Điểm: {round(percentage, 1)}", font=("Times New Roman", 18, "bold"))
    
        # Ghi điểm vào file results.xlsx
        save_results(name, student_class, correct_count, len(questions))

    option_button = [[] for _ in range(4)]  # Danh sách lưu các nút lựa chọn

    # Hiển thị từng câu hỏi và các lựa chọn
    for i, question in enumerate(questions):
        question_text = question['question']
        options = question['options']

        tk.Label(frame, text=f"Câu {i+1}: {question_text}", font=("Times_New_Roman", 12), bg="#F0F0F0").pack(anchor='w', padx=20, pady=5)

        # Biến lưu đáp án của người dùng
        answer_var = tk.StringVar(value="")
        answers_vars.append(answer_var)

        # Hiển thị các lựa chọn dạng RadioButton
        for idx, option in enumerate(options):
            rb = tk.Radiobutton(frame, text=f"{chr(97 + idx)}). {option}", variable=answer_var, value=chr(97 + idx), font=("Times_New_Roman", 12), bg="#F0F0F0")
            rb.pack(anchor='w', padx=40)
            option_button[idx].append(rb)  # Thêm nút vào danh sách

    # Nút nộp bài
    tk.Button(quiz_window, text="Nộp bài", command=submit_answers, font=("Times_New_Roman", 16, "bold"), bg="#4CAF50", fg="white", activebackground="#45a049").pack(pady=10, anchor='center')
    # Nút thoát
    tk.Button(quiz_window, text="Thoát", command=quiz_window.destroy, font=("Times New Roman", 16, "bold"), bg="#800000", fg="white").pack(pady=10, anchor='center')

# Hàm lưu kết quả vào file Excel
def save_results(name, student_class, correct_count, total_questions):
    file_name = "results.xlsx"
    
    # Nếu file đã tồn tại, mở nó, nếu chưa thì tạo mới
    if os.path.exists(file_name):
        workbook = openpyxl.load_workbook(file_name)
    else:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Results"
        # Thêm tiêu đề cho file
        sheet.append(["Họ tên", "Lớp", "Số câu đúng", "Tổng số câu", "Tỷ lệ (%)"])

    # Ghi kết quả vào sheet "Results"
    sheet = workbook["Results"]
    percentage = (correct_count / total_questions) * 100
    sheet.append([name, student_class, correct_count, total_questions, round(percentage, 2)])

    # Lưu file lại
    workbook.save(file_name)
    # Không tắt cửa sổ bài làm sau khi lưu kết quả
    messagebox.showinfo("Thông báo", "Kết quả đã được lưu thành công!")

# Hàm tải danh sách đơn vị bài từ file
def load_units(file_name, sheet_name):
    if os.path.exists(file_name):
        try:
            workbook = openpyxl.load_workbook(file_name)
            sheet = workbook[sheet_name]
            units = set()  # Sử dụng set để loại bỏ trùng lặp
            for row in list(sheet.iter_rows(values_only=True))[1:]:  # Bỏ qua dòng tiêu đề
                if len(row) >= 7:  # Kiểm tra nếu hàng có đủ 7 cột
                    unit = str(row[6]).strip()
                    units.add(unit)
            return sorted([f"Bài {unit}" for unit in units])  # Sắp xếp và định dạng cho combobox
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file đơn vị: {str(e)}")
            return []
    else:
        messagebox.showerror("Lỗi", f"File {file_name} không tồn tại.")
        return []

# Hàm khởi tạo cửa sổ nhập thông tin học sinh
def student_menu():
    exercise_frame = tk.Frame(root, bg="#e6f2ff")  # Khởi tạo test_paper_frame
    exercise_frame.pack(fill="both", expand=True)
    tk.Label(exercise_frame, text="BÀI TẬP TRẮC NGHIỆM NHIỀU LỰA CHỌN", font=("Times New Roman", 18, "bold"), bg="#e6f2ff", fg="#003366").pack(pady=10)
    # Căn giữa các thành phần
    tk.Label(exercise_frame, text="Họ tên học sinh:", font=("Times_New_Roman", 14, "bold"), bg="#e6f2ff", fg="#0000FF").pack(pady=5)
    global name_entry
    name_entry = tk.Entry(exercise_frame, width=30, font=("Times_New_Roman", 14, "bold"))
    name_entry.pack(pady=5)

    tk.Label(exercise_frame, text="Lớp:", font=("Times_New_Roman", 14, "bold"), bg="#e6f2ff", fg="#0000FF").pack(pady=5)
    global class_entry
    class_entry = tk.Entry(exercise_frame, width=30, font=("Times_New_Roman", 14, "bold"))
    class_entry.pack(pady=5)

    tk.Label(exercise_frame, text="Chọn đơn vị bài:", font=("Times_New_Roman", 14, "bold"), bg="#e6f2ff", fg="#0000FF").pack(pady=5)
    global unit_var
    unit_var = tk.StringVar()

    # Lấy danh sách các đơn vị bài từ file
    units = load_units("questions.xlsx", "Questions")
    
    # Tạo combobox để chọn đơn vị bài
    unit_combobox = ttk.Combobox(exercise_frame, textvariable=unit_var, values=units, font=("Times_New_Roman", 14, "bold"))
    unit_combobox.pack(pady=5)
    if units:
        unit_combobox.current(0)  # Đặt đơn vị đầu tiên làm mặc định

    # Nút Bắt đầu bài tập
    tk.Button(exercise_frame, text="Làm bài", command=check_student_info, font=("Times_New_Roman", 16, "bold"), bg="#4CAF50", fg="white", activebackground="#45a049").pack(pady=5)

    # Nút Thoát để đóng cửa sổ
    tk.Button(exercise_frame, text="Thoát", command=exercise_frame.destroy, font=("Times_New_Roman", 16, "bold"), bg="#f44336", fg="white", activebackground="#e53935").pack(pady=5)

# BÀI TẬP TRẮC NGHIỆM ĐÚNG/SAI
# Hàm lấy danh sách các bài từ file Excel
def get_units_from_file(file_name, sheet_name):
    if os.path.exists(file_name):
        try:
            workbook = openpyxl.load_workbook(file_name)
            sheet = workbook[sheet_name]
            units = set()  # Sử dụng set để loại bỏ trùng lặp
            for row in list(sheet.iter_rows(values_only=True))[1:]:  # Bỏ qua dòng tiêu đề
                if len(row) >= 10:  # Kiểm tra nếu hàng có đủ 7 cột
                    unit = str(row[9]).strip()
                    units.add(unit)
            return sorted([f"Bài {unit}" for unit in units])  # Sắp xếp và định dạng cho combobox
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file đơn vị: {str(e)}")
            return []
    else:
        messagebox.showerror("Lỗi", f"File {file_name} không tồn tại.")
        return []

# Hàm kiểm tra thông tin nhập vào (Họ tên, Lớp) trước khi làm bài
def check_student_info_tf():
    name = name_entry.get().strip()
    student_class = class_entry.get().strip()
    selected_unit = unit_var.get()

    if not name or not student_class or not selected_unit:
        messagebox.showerror("Lỗi", "Vui lòng nhập đầy đủ Họ tên, Lớp và Đơn vị bài!")
    else:
        start_quiz_tf(name, student_class, selected_unit)

# Hàm xử lý bắt đầu bài tập trắc nghiệm
def start_quiz_tf(name, student_class, selected_unit):
    questions = []
    answers_vars = []

    # Hàm tải câu hỏi từ file Excel
    def load_true_false_questions(file_name, sheet_name, selected_unit):
        if os.path.exists(file_name):
            try:
                workbook = openpyxl.load_workbook(file_name)
                sheet = workbook[sheet_name]
                for row in list(sheet.iter_rows(values_only=True))[1:]:  # Bỏ qua dòng tiêu đề
                    if len(row) >= 10:  # Kiểm tra nếu hàng có đủ 10 cột ()
                        unit = str(row[9]).strip()
                        if f"Bài {unit}" == selected_unit:  # Chỉ thêm câu hỏi của đơn vị bài đã chọn
                            questions.append({
                                'question': row[0],
                                'options': [str(row[1]).strip(), str(row[3]).strip(), str(row[5]).strip(), str(row[7]).strip()],
                                'correct_answers': [str(row[2]).strip().lower(), str(row[4]).strip().lower(), str(row[6]).strip().lower(), str(row[8]).strip().lower()],
                            })
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể đọc file câu hỏi: {str(e)}")
                return False
        else:
            messagebox.showerror("Lỗi", f"File {file_name} không tồn tại.")
            return False
        return True

    # Tải câu hỏi từ file `questions_true_false.xlsx`
    if not load_true_false_questions("questions_true_false.xlsx", "Questions", selected_unit):
        return

    if not questions:
        messagebox.showerror("Lỗi", "Không có câu hỏi nào cho đơn vị bài đã chọn.")
        return

    # Tạo cửa sổ bài kiểm tra
    quiz_window = tk.Toplevel(root)
    quiz_window.title("Bài kiểm tra trắc nghiệm")
    quiz_window.config(bg="#f0f0f0")
    quiz_window.geometry("950x550")
    quiz_window.transient(root)  # Cửa sổ mới là cửa sổ con của root
    quiz_window.grab_set()  # Đảm bảo chỉ có thể thao tác với cửa sổ này
    # Căn giữa cửa sổ
    center_window(950, 550, quiz_window)
    # Trật tự cửa sổ
    window_order(quiz_window)
    # Tạo canvas cho phép thêm thanh cuộn
    canvas = tk.Canvas(quiz_window, bg="#f0f0f0")
    canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    scrollbar = ttk.Scrollbar(quiz_window, orient="vertical", command=canvas.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    canvas.configure(yscrollcommand=scrollbar.set)
    canvas.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

    frame = tk.Frame(canvas, bg="#f0f0f0")
    canvas.create_window((0, 0), window=frame, anchor="nw")
    tk.Label(frame, text="THÔNG TIN HỌC SINH", font=("Times New Roman", 18, "bold"), fg="#003366").pack(pady=10, anchor="center")
    # Căn giữa thông tin họ tên, lớp, và đề kiểm tra
    tk.Label(frame, text=f"Họ tên: {name}", font=("Times_New_Roman", 14, "bold"), bg="#f0f0f0", fg="#003366").pack(pady=10, anchor="center")
    tk.Label(frame, text=f"Lớp: {student_class}", font=("Times_New_Roman", 14, "bold"), bg="#f0f0f0", fg="#003366").pack(pady=10, anchor="center")
    tk.Label(frame, text=f"Bài tập trắc nghiệm đúng/sai: {selected_unit}", font=("Times_New_Roman", 16, "bold"), bg="#F0F0F0", fg="#0000FF").pack(pady=10, anchor="center")
    # Nhãn để hiển thị điểm
    score_label = tk.Label(frame, text="Số câu đúng: 0 / Tổng số: 0", font=("Times New Roman", 16, "bold"), bg="#f0f0f0", fg="#0000FF")
    score_label.pack(pady=(10, 0), anchor='w')

    total_score_label = tk.Label(frame, text="Tổng điểm: 0", font=("Times New Roman", 16, "bold"), bg="#f0f0f0", fg="#0000FF")
    total_score_label.pack(pady=(10, 0), anchor='w')
    tk.Label(frame, text="PHẦN II. Câu hỏi trắc nghiệm đúng sai.", font=("Times New Roman", 16, "bold"), wraplength=700, justify='left', bg="#f0f0f0", fg="#003366").pack(pady=10)
    def submit_answers_tf():
        correct_count = 0
        total_choices = len(questions) * 4  # Tổng số lựa chọn

        for i, question in enumerate(questions):
            correct_answers = question['correct_answers']  # Đáp án đúng cho các lựa chọn
            user_answers = [answers_vars[i][0].get(), answers_vars[i][1].get(), answers_vars[i][2].get(), answers_vars[i][3].get()]

            # So sánh từng lựa chọn và tính điểm
            for idx, (user_answer, correct_answer) in enumerate(zip(user_answers, correct_answers)):
                option_label = option_labels[i][idx]
                if user_answer == correct_answer:
                    correct_count += 1
                    option_label.config(fg="green")  # Đáp án đúng
                else:
                    option_label.config(fg="red")  # Đáp án sai

        score_label.config(text=f"Số ý trả đúng: {correct_count} / {total_choices}")
        total_score = (correct_count / total_choices) * 10
        total_score_label.config(text=f"Tổng điểm: {total_score:.2f}")

        save_results_tf(name, student_class, correct_count, total_choices)  # Lưu kết quả ở đây

    option_labels = []
    for i, question in enumerate(questions):
        question_text = question['question']
        options = question['options']

        tk.Label(frame, text=f"Câu {i+1}: {question_text}", font=("Times_New_Roman", 16), bg="#f0f0f0", fg="#333333", wraplength=700).pack(anchor='w', padx=20, pady=5)

        answer_var_a = tk.StringVar(value="")
        answer_var_b = tk.StringVar(value="")
        answer_var_c = tk.StringVar(value="")
        answer_var_d = tk.StringVar(value="")
        answers_vars.append([answer_var_a, answer_var_b, answer_var_c, answer_var_d])

        labels = []

        for idx, option in enumerate(options):
            frame_option = tk.Frame(frame, bg="#f0f0f0")
            frame_option.pack(anchor='w', padx=40, pady=2)

            label = tk.Label(frame_option, text=f"{chr(97+idx)}). {option}", font=("Times_New_Roman", 16), bg="#f0f0f0", fg="#333333", wraplength=650)
            label.pack(side=tk.LEFT)
            labels.append(label)

            tk.Radiobutton(frame_option, text="Đúng", variable=answers_vars[i][idx], value="đúng", bg="#f0f0f0", fg="#333333").pack(side=tk.LEFT)
            tk.Radiobutton(frame_option, text="Sai", variable=answers_vars[i][idx], value="sai", bg="#f0f0f0", fg="#333333").pack(side=tk.LEFT)

        option_labels.append(labels)

    # Nút nộp bài
    tk.Button(quiz_window, text="Nộp bài", font=("Times New Roman", 16, "bold"), bg="#4CAF50", fg="white", command=submit_answers_tf).pack(pady=20, anchor='center')
     # Nút thoát
    tk.Button(quiz_window, text="Thoát", command=quiz_window.destroy, font=("Times New Roman", 16, "bold"), bg="#800000", fg="white").pack(pady=10, anchor='center')

## Hàm lưu kết quả vào file Excel
def save_results_tf(name, student_class, correct_count, total_choices):
    file_name = "results.xlsx"
    
    # Nếu file đã tồn tại, mở nó, nếu chưa thì tạo mới
    if os.path.exists(file_name):
        workbook = openpyxl.load_workbook(file_name)
    else:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Results"
        # Thêm tiêu đề cho file
        sheet.append(["Họ tên", "Lớp", "Số lựa chọn đúng", "Tổng số lựa chọn", "Tỷ lệ (%)"])

    # Ghi kết quả vào sheet "Results"
    sheet = workbook["Results"]
    percentage = (correct_count / total_choices) * 100
    sheet.append([name, student_class, correct_count, total_choices, round(percentage, 2)])

    # Lưu file lại
    workbook.save(file_name)
    messagebox.showinfo("Thông báo", "Kết quả đã được lưu thành công!")

# Hàm khởi tạo cửa sổ nhập thông tin học sinh
def student_menu_true_false():
    exercise_frame = tk.Frame(root, bg="#e6f2ff")  # Khởi tạo test_paper_frame
    exercise_frame.pack(fill="both", expand=True)
    tk.Label(exercise_frame, text="BÀI TẬP TRẮC NGHIỆM ĐÚNG/SAI", font=("Times New Roman", 18, "bold"), bg="#e6f2ff", fg="#003366").pack(pady=10)
    # Căn giữa các thành phần
    tk.Label(exercise_frame, text="Họ tên học sinh:", font=("Times_New_Roman", 14, "bold"), bg="#e6f2ff", fg="#0000FF").pack(pady=5)
    global name_entry
    name_entry = tk.Entry(exercise_frame, width=30, font=("Times_New_Roman", 14, "bold"))
    name_entry.pack(pady=5)

    tk.Label(exercise_frame, text="Lớp:", font=("Times_New_Roman", 14, "bold"), bg="#e6f2ff", fg="#0000FF").pack(pady=5)
    global class_entry
    class_entry = tk.Entry(exercise_frame, width=30, font=("Times_New_Roman", 14, "bold"))
    class_entry.pack(pady=5)

    unit_label = tk.Label(exercise_frame, text="Chọn đơn vị bài:", font=("Times New Roman", 14, "bold"), bg="#e6f2ff", fg="#0000FF")
    unit_label.pack(pady=5)
    global unit_var

    unit_var = tk.StringVar()
    units = get_units_from_file("questions_true_false.xlsx", "Questions")  # Lấy danh sách các bài từ file Excel
    unit_combobox = ttk.Combobox(exercise_frame, textvariable=unit_var, font=("Times New Roman", 14, "bold"))
    unit_combobox['values'] = units
    unit_combobox.pack(pady=5)

    # Nút Bắt đầu bài tập
    tk.Button(exercise_frame, text="Làm bài", command=check_student_info_tf, font=("Times_New_Roman", 14, "bold"), bg="#4CAF50", fg="white", activebackground="#45a049").pack(pady=5)

    # Nút Thoát để đóng cửa sổ
    tk.Button(exercise_frame, text="Thoát", command=exercise_frame.destroy, font=("Times_New_Roman", 14, "bold"), bg="#f44336", fg="white", activebackground="#e53935").pack(pady=5)

# KIỂM TRA THỬ
def check_student_info_exercise():
    name = name_entry.get().strip()
    student_class = class_entry.get().strip()

    if not name or not student_class:
        messagebox.showerror("Lỗi", "Vui lòng nhập đầy đủ Họ tên và Lớp!")
    else:
        selected_unit = unit_var.get()  # Lấy đơn vị bài đã chọn
        start_combined_quiz_exercise(name, student_class, selected_unit)

# Hàm xử lý bắt đầu bài kiểm tra cho cả hai phần
def start_combined_quiz_exercise(name, student_class, selected_unit):
    multiple_choice_questions = []
    true_false_questions = []
    answers_vars_mc = []
    answers_vars_tf = []

    # Hàm tải câu hỏi trắc nghiệm 4 lựa chọn từ file Excel
    def load_multiple_choice_questions(file_name, sheet_name, selected_unit):
        if os.path.exists(file_name):
            try:
                workbook = openpyxl.load_workbook(file_name)
                sheet = workbook[sheet_name]
                for row in list(sheet.iter_rows(values_only=True))[1:]:
                    if len(row) >= 7:
                        question = row[0]
                        options = [str(row[1]).strip(), str(row[2]).strip(), str(row[3]).strip(), str(row[4]).strip()]
                        correct_answer = str(row[5]).strip().lower()
                        unit = str(row[6]).strip()
                        
                        if f"Bài {unit}" == selected_unit:
                            multiple_choice_questions.append({
                                'question': question,
                                'options': options,
                                'correct_answer': correct_answer,
                            })
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể đọc file câu hỏi: {str(e)}")
                return False
        else:
            messagebox.showerror("Lỗi", f"File {file_name} không tồn tại.")
            return False
        return True

    # Tải câu hỏi từ file `questions.xlsx`
    if not load_multiple_choice_questions("questions.xlsx", "Questions", selected_unit):
        return

    # Hàm tải câu hỏi đúng/sai theo mẫu từ file Excel
    def load_true_false_questions(file_name, sheet_name, selected_unit):
        if os.path.exists(file_name):
            try:
                workbook = openpyxl.load_workbook(file_name)
                sheet = workbook[sheet_name]
                for row in list(sheet.iter_rows(values_only=True))[1:]:
                    if len(row) >= 10:
                        unit = str(row[9]).strip()
                        if f"Bài {unit}" == selected_unit:
                            true_false_questions.append({
                                'question': row[0],
                                'options': [str(row[1]).strip(), str(row[3]).strip(), str(row[5]).strip(), str(row[7]).strip()],
                                'correct_answers': [str(row[2]).strip().lower(), str(row[4]).strip().lower(), str(row[6]).strip().lower(), str(row[8]).strip().lower()],
                            })
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể đọc file câu hỏi: {str(e)}")
                return False
        else:
            messagebox.showerror("Lỗi", f"File {file_name} không tồn tại.")
            return False
        return True

    if not load_true_false_questions("questions_true_false.xlsx", "Questions", selected_unit):
        return

    if not multiple_choice_questions and not true_false_questions:
        messagebox.showerror("Lỗi", "Không có câu hỏi nào cho đơn vị bài đã chọn.")
        return

    random.shuffle(multiple_choice_questions)
    random.shuffle(true_false_questions)

    quiz_window = tk.Toplevel(root)
    quiz_window.title("Bài kiểm tra trắc nghiệm")
    quiz_window.geometry("800x550")
    quiz_window.configure(bg="#f0f0f0")
    # Căn giữa cửa sổ
    center_window(800, 550, quiz_window)
    # Trật tự cửa sổ
    window_order(quiz_window)
    
    canvas = tk.Canvas(quiz_window, bg="#f0f0f0")
    canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    scrollbar = ttk.Scrollbar(quiz_window, orient="vertical", command=canvas.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    canvas.configure(yscrollcommand=scrollbar.set)
    canvas.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

    frame = tk.Frame(canvas, bg="#f0f0f0")
    canvas.create_window((0, 0), window=frame, anchor="nw")

    title_frame = tk.Frame(frame, bg="#f0f0f0")
    title_frame.pack(pady=10)
    
    title_label = tk.Label(title_frame, text="BÀI KIỂM TRA THỬ", font=("Times New Roman", 20, "bold"), fg="#003366", bg="#f0f0f0")
    title_label.pack()

    name_label = tk.Label(title_frame, text=f"Họ tên: {name}", font=("Times New Roman", 14, "bold"), bg="#f0f0f0", fg="#0000FF")
    name_label.pack(anchor='w')

    class_label = tk.Label(title_frame, text=f"Lớp: {student_class}", font=("Times New Roman", 14, "bold"), bg="#f0f0f0", fg="#0000FF")
    class_label.pack(anchor='w')
   
    # Label hiển thị số câu đúng và điểm
    result_label = tk.Label(frame, text="Số câu đúng: 0/0| Điểm: 0", font=("Times New Roman", 16, "bold"), bg="#f0f0f0", fg="#0000FF")
    result_label.pack(anchor='w', padx=20, pady=10)

    # Phần I: Trắc nghiệm nhiều lựa chọn
    tk.Label(frame, text="PHẦN I. Câu trắc nghiệm nhiều lựa chọn:", font=("Times New Roman", 20, "bold"), wraplength=700, justify='left', bg="#f0f0f0", fg="#003366").pack(pady=10, anchor='center')

    for i, question in enumerate(multiple_choice_questions):
        tk.Label(frame, text=f"Câu {i+1}: {question['question']}", font=("Times New Roman", 16), wraplength=600, justify='left', bg="#f0f0f0", fg="#333333").pack(anchor='w', padx=20, pady=5)

        answer_var = tk.StringVar(value="")
        answers_vars_mc.append(answer_var)

        for idx, option in enumerate(question['options']):
            tk.Radiobutton(frame, text=f"{chr(97+idx)}). {option}", variable=answer_var, value=chr(97+idx), wraplength=550, justify='left', bg="#f0f0f0", fg="#333333", font=("Times New Roman", 16)).pack(anchor='w', padx=40)

    # Phần II: Trắc nghiệm đúng/sai
    tk.Label(frame, text="PHẦN II. Câu hỏi đúng/sai (cho từng lựa chọn):", font=("Times New Roman", 20, "bold"), wraplength=700, justify='left', bg="#f0f0f0", fg="#003366").pack(pady=10, anchor='center')

    for i, question in enumerate(true_false_questions):
        tk.Label(frame, text=f"Câu {i+1}: {question['question']}", font=("Times New Roman", 16), wraplength=600, justify='left', bg="#f0f0f0", fg="#333333").pack(anchor='w', padx=20, pady=5)

        answers_vars_tf.append([tk.StringVar(value="") for _ in range(len(question['options']))])

        for idx, option in enumerate(question['options']):
            tk.Label(frame, text=f"{chr(97+idx)}). {option}", font=("Times New Roman", 16), wraplength=600, justify='left', bg="#f0f0f0", fg="#333333").pack(anchor='w', padx=20)
            tk.Radiobutton(frame, text="Đúng", variable=answers_vars_tf[i][idx], value="đúng", wraplength=550, justify='left', bg="#f0f0f0", fg="#333333", font=("Times New Roman", 16)).pack(anchor='w', padx=60)
            tk.Radiobutton(frame, text="Sai", variable=answers_vars_tf[i][idx], value="sai", wraplength=550, justify='left', bg="#f0f0f0", fg="#333333", font=("Times New Roman", 16)).pack(anchor='w', padx=60)

    def save_results_to_excel(name, student_class, score_mc, score_tf, total_questions, score):
        results_file = "results.xlsx"

        # Nếu file không tồn tại, tạo mới
        if not os.path.exists(results_file):
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = "Results"
            sheet.append(["Họ tên", "Lớp", "Số câu đúng", "Tổng số câu", "Điểm"])
        else:
            workbook = openpyxl.load_workbook(results_file)
            sheet = workbook["Results"]

        # Ghi thông tin kết quả
        sheet.append([name, student_class, score_mc + score_tf, total_questions, score])
        workbook.save(results_file)

    def check_answers():
        score_mc = 0
        score_tf = 0
        total_mc = len(multiple_choice_questions)
        total_tf = len(true_false_questions)

        for i, question in enumerate(multiple_choice_questions):
            selected_answer = answers_vars_mc[i].get()
            if selected_answer and selected_answer == question['correct_answer']:
                score_mc += 1

        for i, question in enumerate(true_false_questions):
            for idx, option in enumerate(question['options']):
                selected_answer = answers_vars_tf[i][idx].get()
                if selected_answer == "đúng" and question['correct_answers'][idx] == "đúng":
                    score_tf += 1
                elif selected_answer == "sai" and question['correct_answers'][idx] == "sai":
                    score_tf += 1

        total_score = score_mc + score_tf
        total_questions = len(multiple_choice_questions) + sum([len(q['options']) for q in true_false_questions])

        # Tính điểm theo thang 10
        score = round((total_score/ total_questions) * 10, 2)

        result_label.config(text=f"Số câu đúng: {score_mc + score_tf}/{total_questions} | Điểm: {score:.1f}")
        save_results_to_excel(name, student_class, score_mc, score_tf, total_questions, score)

    submit_button = tk.Button(quiz_window, text="Nộp bài", command=check_answers, font=("Times New Roman", 16), bg="#4CAF50", fg="white")
    submit_button.pack(pady=20)
     # Nút thoát
    tk.Button(quiz_window, text="Thoát", command=quiz_window.destroy, font=("Times New Roman", 16), bg="#800000", fg="white").pack(pady=10, anchor='center')
def student_menu_exercise():
    exercise_frame = tk.Frame(root, bg="#e6f2ff")  # Khởi tạo test_paper_frame
    exercise_frame.pack(fill="both", expand=True)
    # Đảm bảo `unit_var` là toàn cục
    global unit_var
    tk.Label(exercise_frame, text="BÀI KIỂM TRA THỬ", font=("Times New Roman", 18, "bold"), bg="#e6f2ff", fg="#003366").pack(pady=10)
    # Căn giữa các thành phần
    tk.Label(exercise_frame, text="Họ tên học sinh:", font=("Times_New_Roman", 14, "bold"), bg="#e6f2ff", fg="#0000FF").pack(pady=5)
    global name_entry
    name_entry = tk.Entry(exercise_frame, width=30, font=("Times_New_Roman", 14, "bold"))
    name_entry.pack(pady=5)

    tk.Label(exercise_frame, text="Lớp:", font=("Times_New_Roman", 14, "bold"), bg="#e6f2ff", fg="#0000FF").pack(pady=5)
    global class_entry
    class_entry = tk.Entry(exercise_frame, width=30, font=("Times_New_Roman", 14, "bold"))
    class_entry.pack(pady=5)

    # ComboBox chọn đơn vị bài
    tk.Label(exercise_frame, text="Chọn Đơn vị bài:", font=("Times New Roman", 14, "bold"), bg="#e6f2ff", fg="#0000FF").pack(pady=10)
    unit_var = tk.StringVar()
    unit_combobox = ttk.Combobox(exercise_frame, textvariable=unit_var, font=("Times New Roman", 14))
    unit_combobox['values'] = [f"Bài {i}" for i in range(1, 33)]  # Giả định có 32 đơn vị bài
    unit_combobox.pack(pady=10)

    # Nút Bắt đầu bài tập
    tk.Button(exercise_frame, text="Làm bài", command=check_student_info_exercise, font=("Times_New_Roman", 14, "bold"), bg="#4CAF50", fg="white", activebackground="#45a049").pack(pady=5)

    # Nút Thoát để đóng cửa sổ
    tk.Button(exercise_frame, text="Thoát", command=exercise_frame.destroy, font=("Times_New_Roman", 14, "bold"), bg="#f44336", fg="white", activebackground="#e53935").pack(pady=5)


# BÀI KIỂM TRA
# Hàm kiểm tra thông tin nhập vào (Họ tên, Lớp)
def check_student_info_kt():
    name = name_entry.get().strip()  # Sử dụng strip() thay vì trim()
    student_class = class_entry.get().strip()  # Sử dụng strip() thay vì trim()

    if not name or not student_class:
        messagebox.showerror("Lỗi", "Vui lòng nhập đầy đủ Họ tên và Lớp!")
    else:
        start_combined_quiz(name, student_class)

# Hàm xử lý bắt đầu bài kiểm tra cho cả hai phần
def start_combined_quiz(name, student_class):
    multiple_choice_questions = []
    true_false_questions = []
    answers_vars_mc = []
    answers_vars_tf = []

    # Hàm tải câu hỏi trắc nghiệm 4 lựa chọn từ file Excel
    def load_multiple_choice_questions(file_name, sheet_name):
        if os.path.exists(file_name):
            try:
                workbook = openpyxl.load_workbook(file_name)
                sheet = workbook[sheet_name]
                for row in list(sheet.iter_rows(values_only=True))[1:]:
                    if len(row) >= 6:
                        multiple_choice_questions.append({
                            'question': row[0],
                            'options': [str(row[1]).strip(), str(row[2]).strip(), str(row[3]).strip(), str(row[4]).strip()],
                            'correct_answer': str(row[5]).strip().lower(),
                        })
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể đọc file câu hỏi: {str(e)}")
                return False
        else:
            messagebox.showerror("Lỗi", f"File {file_name} không tồn tại.")
            return False
        return True

    # Hàm tải câu hỏi đúng/sai theo mẫu từ file Excel
    def load_true_false_questions(file_name, sheet_name):
        if os.path.exists(file_name):
            try:
                workbook = openpyxl.load_workbook(file_name)
                sheet = workbook[sheet_name]
                for row in list(sheet.iter_rows(values_only=True))[1:]:
                    if len(row) >= 9:
                        true_false_questions.append({
                            'question': row[0],
                            'options': [str(row[1]).strip(), str(row[3]).strip(), str(row[5]).strip(), str(row[7]).strip()],
                            'correct_answers': [str(row[2]).strip().lower(), str(row[4]).strip().lower(), str(row[6]).strip().lower(), str(row[8]).strip().lower()],
                        })
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể đọc file câu hỏi: {str(e)}")
                return False
        else:
            messagebox.showerror("Lỗi", f"File {file_name} không tồn tại.")
            return False
        return True

    # Tải câu hỏi từ file Excel
    if not load_multiple_choice_questions("questions.xlsx", "Questions") or not load_true_false_questions("questions_true_false.xlsx", "Questions"):
        return

    if not multiple_choice_questions and not true_false_questions:
        messagebox.showerror("Lỗi", "Chưa có câu hỏi nào được nhập hoặc định dạng không đúng.")
        return

    # Trộn ngẫu nhiên các câu hỏi
    random.shuffle(multiple_choice_questions)
    random.shuffle(true_false_questions)

    # Tạo cửa sổ bài kiểm tra
    quiz_window = tk.Toplevel(root)
    quiz_window.title("Bài kiểm tra trắc nghiệm")
    quiz_window.geometry("800x550")
    quiz_window.configure(bg="#f0f0f0")
    # Căn giữa cửa sổ
    center_window(800, 550, quiz_window)
    # Trật tự cửa sổ
    window_order(quiz_window)
    
    canvas = tk.Canvas(quiz_window, bg="#f0f0f0")
    canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    scrollbar = ttk.Scrollbar(quiz_window, orient="vertical", command=canvas.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    canvas.configure(yscrollcommand=scrollbar.set)
    canvas.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

    frame = tk.Frame(canvas, bg="#f0f0f0")
    canvas.create_window((0, 0), window=frame, anchor="nw")

    # Hiển thị tiêu đề bài kiểm tra
    title_frame = tk.Frame(frame, bg="#f0f0f0")
    title_frame.pack(pady=10)

    title_label = tk.Label(title_frame, text="Bài kiểm tra trắc nghiệm", font=("Times New Roman", 24, "bold"), fg="#003366", bg="#f0f0f0")
    title_label.pack()

    name_label = tk.Label(title_frame, text=f"Họ tên: {name}", font=("Times New Roman", 16), bg="#f0f0f0", fg="#003366")
    name_label.pack(anchor='w')

    class_label = tk.Label(title_frame, text=f"Lớp: {student_class}", font=("Times New Roman", 16), bg="#f0f0f0", fg="#003366")
    class_label.pack(anchor='w')

    # Phần I: Trắc nghiệm nhiều lựa chọn
    tk.Label(frame, text="PHẦN I. Câu trắc nghiệm nhiều lựa chọn:", font=("Times New Roman", 20), wraplength=700, justify='left', bg="#f0f0f0", fg="#003366").pack(pady=10, anchor='center')

    for i, question in enumerate(multiple_choice_questions):
        tk.Label(frame, text=f"Câu {i+1}: {question['question']}", font=("Times New Roman", 16), wraplength=600, justify='left', bg="#f0f0f0", fg="#333333").pack(anchor='w', padx=20, pady=5)

        answer_var = tk.StringVar(value="")
        answers_vars_mc.append(answer_var)

        # Đáp án không trộn
        options = question['options']

        for idx, option in enumerate(options):
            tk.Radiobutton(frame, text=f"{chr(97+idx)}). {option}", variable=answer_var, value=chr(97+idx), wraplength=550, justify='left', bg="#f0f0f0", fg="#333333", font=("Times New Roman", 16)).pack(anchor='w', padx=40)

    # Phần II: Trắc nghiệm đúng/sai
    tk.Label(frame, text="PHẦN II. Câu hỏi đúng/sai (cho từng lựa chọn):", font=("Times New Roman", 20), wraplength=700, justify='left', bg="#f0f0f0", fg="#003366").pack(pady=10, anchor='center')

    for i, question in enumerate(true_false_questions):
        tk.Label(frame, text=f"Câu {i+1}: {question['question']}", font=("Times New Roman", 16), wraplength=600, justify='left', bg="#f0f0f0", fg="#333333").pack(anchor='w', padx=20, pady=5)

        answers_vars_tf.append([tk.StringVar(value="") for _ in range(len(question['options']))])

        # Đáp án không trộn
        options = question['options']

        for idx, option in enumerate(options):
            tk.Label(frame, text=f"{chr(97+idx)}). {option}", font=("Times New Roman", 16), wraplength=600, justify='left', bg="#f0f0f0", fg="#333333").pack(anchor='w', padx=40)
            tk.Radiobutton(frame, text="Đúng", variable=answers_vars_tf[i][idx], value="đúng", wraplength=550, justify='left', bg="#f0f0f0", fg="#333333", font=("Times New Roman", 16)).pack(anchor='w', padx=60)
            tk.Radiobutton(frame, text="Sai", variable=answers_vars_tf[i][idx], value="sai", wraplength=550, justify='left', bg="#f0f0f0", fg="#333333", font=("Times New Roman", 16)).pack(anchor='w', padx=60)
    # Hàm lưu kết quả vào file results.xlsx
    def save_results_to_excel(name, student_class, score_mc, score_tf, total_questions, score):
        results_file = "results.xlsx"

        # Nếu file không tồn tại, tạo mới
        if not os.path.exists(results_file):
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = "Results"
            sheet.append(["Họ tên", "Lớp", "Số câu đúng", "Tổng số câu", "Điểm"])
        else:
            workbook = openpyxl.load_workbook(results_file)
            sheet = workbook["Results"]

        # Ghi kết quả vào file
        sheet.append([name, student_class, score_mc + score_tf, total_questions, score])
        workbook.save(results_file)
    # Nút nộp bài
    def submit_answers():
        # Kiểm tra nếu có câu hỏi nào chưa được trả lời
        for i, answer_var in enumerate(answers_vars_mc):
            if answer_var.get() == "":
                messagebox.showerror("Lỗi", f"Vui lòng chọn đáp án cho câu hỏi {i + 1} trong phần I.")
                return

        for i, answer_vars in enumerate(answers_vars_tf):
            for idx, answer_var in enumerate(answer_vars):
                if answer_var.get() == "":
                    messagebox.showerror("Lỗi", f"Vui lòng chọn đáp án cho câu hỏi {i + 1}, lựa chọn {chr(97+idx)} trong phần II.")
                    return
        
        # Tính điểm phần I
        score_mc = 0
        for i, question in enumerate(multiple_choice_questions):
            if answers_vars_mc[i].get() == question['correct_answer']:
                score_mc += 1

        # Tính điểm phần II
        score_tf = 0
        for i, question in enumerate(true_false_questions):
            for idx, answer_var in enumerate(answers_vars_tf[i]):
                if answer_var.get() == question['correct_answers'][idx]:
                    score_tf += 1

        # Tổng điểm
        total_score = score_mc + score_tf
        total_questions = len(multiple_choice_questions) + sum([len(q['options']) for q in true_false_questions])

        # Tính điểm theo thang 10
        score = round((total_score/ total_questions) * 10, 2)

        # Lưu kết quả vào file
        save_results_to_excel(name, student_class, score_mc, score_tf, total_questions, score)
        # Hiện kết quả
        quiz_window.withdraw()  # Ẩn cửa sổ bài kiểm tra
        result_window = tk.Toplevel(root)
        result_window.title("Kết quả bài kiểm tra")
        result_window.geometry("500x400")
        result_window.configure(bg="#f0f0f0")
        # Căn giữa cửa sổ
        center_window(500, 400, result_window)
        # Trật tự cửa sổ
        window_order(result_window)
        
        result_frame = tk.Frame(result_window, bg="#f0f0f0")
        result_frame.pack(pady=20)

        tk.Label(result_frame, text="KẾT QUẢ BÀI KIỂM TRA", font=("Times New Roman", 20, "bold"), bg="#f0f0f0", fg="#003366").pack(pady=10)

        tk.Label(result_frame, text=f"Họ tên: {name}", font=("Times New Roman", 16, "bold"), bg="#f0f0f0", fg="#0000FF").pack(pady=5)
        tk.Label(result_frame, text=f"Lớp: {student_class}", font=("Times New Roman", 16, "bold"), bg="#f0f0f0", fg="#0000FF").pack(pady=5)
        tk.Label(result_frame, text=f"Số câu đúng: {score_mc + score_tf}/{total_questions}", font=("Times New Roman", 16, "bold"), bg="#f0f0f0", fg="#FF0000").pack(pady=5)
        tk.Label(result_frame, text=f"Điểm: {score}", font=("Times New Roman", 16, "bold"), bg="#f0f0f0", fg="#FF0000").pack(pady=5)

        tk.Button(result_frame, text="Đóng", font=("Times New Roman", 16, "bold"), bg="#4CAF50", fg="#003366", command=result_window.destroy).pack(pady=20)

    tk.Button(quiz_window, text="Nộp bài", command=submit_answers, font=("Times New Roman", 16, "bold"), bg="#4CAF50", fg="white").pack(pady=20)
  
# Hàm khởi tạo cửa sổ nhập thông tin học sinh
def student_menu_test_paper():
    test_paper_frame = tk.Frame(root, bg="#e6f2ff")  # Khởi tạo test_paper_frame
    test_paper_frame.pack(fill="both", expand=True)
    tk.Label(test_paper_frame, text="NHẬP THÔNG TIN HỌC SINH", font=("Times_New_Roman", 18, 'bold'), bg="#e6f2ff", fg="#003366").pack(pady=5)
    tk.Label(test_paper_frame, text="Họ tên học sinh:", font=("Times_New_Roman", 14, 'bold'), bg="#e6f2ff", fg="#0000FF").pack(pady=5)
    global name_entry
    name_entry = tk.Entry(test_paper_frame, width=30, font=("Times_New_Roman", 14, 'bold'))
    name_entry.pack(pady=5)

    tk.Label(test_paper_frame, text="Lớp:", font=("Times_New_Roman", 14, 'bold'), bg="#e6f2ff", fg="#0000FF").pack(pady=5)
    global class_entry
    class_entry = tk.Entry(test_paper_frame, width=30, font=("Times_New_Roman", 14, 'bold'))
    class_entry.pack(pady=5)

   # Nút Bắt đầu bài tập
    tk.Button(test_paper_frame, text="Làm bài", command=check_student_info_kt, font=("Times_New_Roman", 16, 'bold'), bg="#4CAF50", fg="white", activebackground="#45a049").pack(pady=5)

    # Nút Thoát để đóng cửa sổ
    tk.Button(test_paper_frame, text="Thoát", command=test_paper_frame.destroy, font=("Times_New_Roman", 16, 'bold'), bg="#f44336", fg="white", activebackground="#e53935").pack(pady=5)

# MỞ FILE HƯỚNG DẪN: Hàm mở file huongdan.pdf
def open_help():
    help_file = "huongdan.pdf"
    if os.path.exists(help_file):
        subprocess.Popen(['start', help_file], shell=True)  # Mở file trong hệ điều hành
    else:
        messagebox.showerror("Lỗi", "File huongdan.pdf không tồn tại!")
        
#GIAO DIỆN CHÍNH
# Hàm thoát ứng dụng
def exit_app():
    root.destroy()
# Giao diện
quiz = Quiz()
# Tạo cửa sổ chính
root = tk.Tk()
root.title("Ứng dụng trả lời trắc nghiệm")
root.geometry("1400x700")
# Khởi tạo file danh sách tài khoản
create_accounts_file()
# Tạo menu
menubar = tk.Menu(root)
root.option_add('*TearOff', False)
root.option_add('*Menu*Font', ('Times_New_Roman', 16))  # Cỡ chữ lớn cho menu
# 1. Cấp tài khoản (Admin)
admin_accounts = tk.Menu(menubar, tearoff=0)
admin_accounts.add_command(label="Cấp tài khoản", command=add_account)
admin_accounts.add_command(label="Tải danh sách mẫu tài khoản", command=download_sample_file)
admin_accounts.add_command(label="Nạp danh sách tài khoản vào hệ thống", command=import_accounts)
menubar.add_cascade(label="Admin", menu=admin_accounts)

# 2. Tạo menu Đăng nhập hệ thống
account_menu_teacher_student = tk.Menu(menubar, tearoff=0)
account_menu_teacher_student.add_command(label="Đăng nhập", command=lambda: user_login("Teacher"))
account_menu_teacher_student.add_command(label="Thay đổi tài khoản và mật khẩu", command=lambda: change_account("Teacher"))
menubar.add_cascade(label="Đăng Nhập", menu=account_menu_teacher_student)
# 3. Tạo menu Giáo viên
teacher_menu_item = tk.Menu(menubar, tearoff=0)
# Tạo menu con "Nhập câu hỏi"
enter_question_menu = tk.Menu(teacher_menu_item, tearoff=0)
enter_question_menu.add_command(label="Trắc nghiệm nhiều lựa chọn", command=multiple_choice)
enter_question_menu.add_command(label="Trắc nghiệm đúng/sai", command=true_false)
teacher_menu_item.add_cascade(label="Nhập câu hỏi", menu=enter_question_menu)
# Tạo menu còn Trộn đề
mix_menu = tk.Menu(menubar, tearoff=0)
mix_menu.add_command(label="Trộn đề từ ngân hàng câu hỏi", command=open_shuffle_window)
mix_menu.add_command(label="Trộn đề từ file", command=open_shuffle_window_word)
teacher_menu_item.add_cascade(label="Trộn đề kiểm tra", menu=mix_menu )
# Tạo menu còn Bài dạy
lesson_menu = tk.Menu(menubar, tearoff=0)
lesson_menu.add_command(label="Bài tập trắc nghiệm", command=teacher_menu_quiz_lesson)
lesson_menu.add_command(label="Lý thuyết", command=teacher_menu_theory_lesson)
teacher_menu_item.add_cascade(label="Bài dạy", menu=lesson_menu )
# Tạo menu con Kết quả
results_menu = tk.Menu(menubar, tearoff=0)
results_menu.add_command(label="Xem kết quả", command=show_results)
results_menu.add_command(label="Tải bảng kết quả", command=download_results)
teacher_menu_item.add_cascade(label="Kết quả", menu=results_menu )

menubar.add_cascade(label="Giáo viên", menu=teacher_menu_item)
# 4. Tạo menu Học sinh
student_menu_item = tk.Menu(menubar, tearoff=0)
student_menu_item.add_command(label="Bài Tập trắc nghiệm nhiều lựa chọn", command=student_menu)
student_menu_item.add_command(label="Bài tập trắc nghiệm đúng/sai", command=student_menu_true_false)
student_menu_item.add_command(label="Bài kiểm tra thử", command=student_menu_exercise)
menubar.add_cascade(label="Học sinh", menu=student_menu_item)
# 5 Tạo menu kiểm tra
menubar.add_cascade(label="Kiểm tra", command=student_menu_test_paper)
# 6 Tạo menu hướng dẫn sử dụng
menubar.add_cascade(label="Hướng dẫn sử dụng", command=open_help)
# 7. Tạo nút thoát ứng dụng
menubar.add_cascade(label="Thoát", command=exit_app)
# Ẩn các menu "Giáo Viên" và "Học Sinh" khi chưa đăng nhập
hide_menus()
# Gắn menu vào cửa sổ
root.config(menu=menubar)

# Khởi chạy chương trình
root.mainloop()
